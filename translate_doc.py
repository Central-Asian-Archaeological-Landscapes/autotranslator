'''
Created on 31 May 2022

@author: turch
'''
from deep_translator import GoogleTranslator as transl #uses deep_translator instead of googletrans because
# googletrans maxes out times that you can run entries - this has a much higher threshold
#although if it starts timing out lots can change to MicrosoftTranslator - still effective 
#import openpyxl as xl
import fnmatch as fn
import os
import src.spellcheck as spellcheck #NOTE: If running in IDLE delete 'from src' if running in Eclipse keep it
import tkinter as tk
import docx
from src.resetter import ResetGUI
import re

import pandas as pd
import numpy as np
import pytesseract
import pdf2image
from pdf2image import convert_from_path
import PIL

class TranslateDoc:
    pdf_content = []
    doc = []
    tables = []
    translated_table = []
    row_length = []
    column_length = []
    translated_doc = []
    err_txt = ' '
    counter0_txt = '...'
    ldat_txt = '...'
    counter1_txt = '...'
    counter2_txt = '...'
    ent_txt = '...'
    chk3_txt = '...'
    resett = False
    
    def update(self, err, counter0, ldat, counter1, counter2, ent, chk3, windowx): #function to allow window to continue updating
        err.config(text = self.err_txt)
        counter0.config(text = self.counter0_txt)
        ldat.config(text = self.ldat_txt)
        counter1.config(text = self.counter1_txt)
        counter2.config(text = self.counter2_txt)
        ent.config(text = self.ent_txt)
        chk3.config(text = self.chk3_txt)
        windowx.after(500, self.update, err, counter0, ldat, counter1, counter2, ent, chk3, windowx)
        
    def load_data(self, filename, filetype, ilang):
        
        if filetype == 'Word':
            try:
                untrans = docx.Document(filename)
                self.untrans = untrans
                self.ldat_txt = 'Text extracted from file'
            except Exception as e: 
                windoww = tk.Tk()
                fram = tk.Frame(master=windoww, width = 150, height = 150, padx = 5, pady=5)
                fram.pack(fill = tk.X, side = tk.TOP)
                err = tk.Label(master=fram, text = 'ERROR:' + str(e))
                err.pack(fill=tk.X, side = tk.TOP)
                tip = tk.Label(master=fram, text = 'If Permission Error check permissions on files and tick "allowed to edit" for everyone or check that files are not open')
                tip.pack(fill=tk.X, side = tk.TOP)
                
                framee = tk.Frame(master=windoww, width = 150, height = 150, padx = 5, pady=5)
                framee.pack()
                new_btn = tk.Button(master=framee, text = 'BACK', command = ResetGUI().restart())
                new_btn.pack(fill = tk.X, side = tk.BOTTOM)
                windoww.mainloop()
        elif filetype == 'PDF':
            pn = os.path.split(os.path.dirname(os.path.abspath(__file__)))[0]
            poppler = os.path.join(pn, 'src', 'data', 'poppler-0.68.0', 'bin')

            #get it to read PDF extract text as object translate it and output as a doc
            #read_pdf = PyPDF2.PdfFileReader(filename)
            #page_number = threading.Thread(target=read_pdf.getNumPages()).start
            #print(poppler)
            PIL.Image.MAX_IMAGE_PIXELS = 1500000000000
            
            images = convert_from_path(filename, poppler_path =poppler)
            self.counter0_txt = 'Extracted pages from PDF'
            tess = os.path.join(pn, 'src', 'TesseractOCR', 'tesseract.exe')
            pytesseract.pytesseract.tesseract_cmd = tess
        
        
            if ilang == 'ru':
                lan = 'rus'
            elif ilang == 'en':
                lan = 'eng'
            elif ilang == 'kk':
                lan = 'kaz'
            elif ilang == 'ky':
                lan = 'kir'
            elif ilang == 'fr':
                lan = 'fra'
            elif ilang == 'ar':
                lan = 'ara'
            elif ilang == 'de':
                lan = 'deu'
            elif ilang == 'zh_CN':
                lan = 'chi_sim'
            elif ilang == 'es':
                lan = 'spa'
            elif ilang == 'uz':
                lan = 'uzb'
            elif ilang == 'tg':
                lan = 'tgk'
            elif ilang == 'it':
                lan = 'ita'
            else:
                lan = 'eng'
            n=0
            
            for img in enumerate(images):
                image = img.conver('1')
                txt = pytesseract.image_to_string(image, lang = lan)
                self.pdf_content.append(str(txt))
                print(str(txt))
                n+=1
                self.counter0_txt = str(n) + 'pages extracted from PDF'
            '''
            for page, page_data in enumerate(convert):
                txt = pytesseract.image_to_string(page_data, lang = lan)
                pg_txt =("Page # {} - {}".format(str(page), txt))
                self.pdf_content.append(pg_txt)
                n+=1
                print(str(pg_txt))
                self.counter0_txt = str(n) + ' pages processed from PDF'
            print(str(self.pdf_content) + 'done')
        '''
                        
            self.ldat_txt = 'Text extracted from pages'
        else:
            pass
        
    def check_data(self, filetype, ilang, olang):
        
        spell = spellcheck.Spelling()
        if ilang == 'ru' and olang == 'en':
            gloss = spell.create_glossary(ilang)
        elif ilang == 'zh_CN' and olang == 'en':
            gloss = spell.create_glossary(ilang)
        if olang == 'ru' and ilang == 'en':
            gloss = spell.create_glossary(olang)
        if olang == 'zh_CN' and ilang == 'en':
            gloss = spell.create_glossary(olang)
        else:
            gloss = {'1':['1']}
            
        
        if filetype == 'Word':
                if len(self.untrans.tables) > 0:
                    for table in self.untrans.tables:
                        transl_table = []
                        self.row_length.append(len(table.rows))
                        print(self.row_length)
                        self.column_length.append(len(table.columns))
                        print(self.column_length)
                        for row in table.rows:
                            for cell in row.cells:
                                cell_text = []
                                for para in cell.paragraphs:
                                    parag = str(para.text)
                                    checked_table_parag = self.check_data2(filetype, ilang, olang, parag, gloss)
                                    cell_text.append(checked_table_parag)
                                nu_cell = '\n'.join(cell_text)
                                #itable.append(nu_cell)
                                print(nu_cell)
                                self.tables.append(nu_cell)
                                transl_cell = self.translate_data(ilang, olang, nu_cell)
                                print(transl_cell)
                                transl_table.append(transl_cell)
                        self.translated_table.append(transl_table)
                                
                                    
                else:
                    pass
                
                if len(self.untrans.paragraphs) > 0:
                    n = 0
                    for para in self.untrans.paragraphs:
                        parag = str(para.text)
                        checked_parag = self.check_data2(filetype, ilang, olang, parag, gloss)
                        self.doc.append(checked_parag)
                        n = n + 1
                        self.counter1_txt = str(n) + ' paragraphs checked out of ' + str(len(self.untrans.paragraphs))
                    
                    for parag in self.doc:
                        transl_parag = self.translate_data(ilang, olang, parag)
                        self.translated_doc.append(transl_parag)
                        print(transl_parag)
                        
        elif filetype == 'PDF':
            for page in self.pdf_content:
                nu_page = str(page).replace('\n', ' ')
                #checked_page = self.check_data2(filetype, ilang, olang, str(nu_page), gloss)
                transl_page = self.translate_data(ilang, olang, nu_page)
                print(transl_page)
                self.translated_doc.append(transl_page)
        else:
            pass
        
    def check_data2(self, filetype, ilang, olang, parag, gloss):
        spell = spellcheck.Spelling()
        for word in parag.split(' '): #splits paragraph into individual words
            
            punc_free = re.sub(r'(?s:[,.!:;"?])\Z', '', word) #removes punctuation from word 
            print(punc_free)
            #print(punc_free)
            if len(str(punc_free)) == 0 or str(punc_free) == None:
                continue
            elif len(str(punc_free)) > 1 and str(punc_free).strip()[0].isupper(): #if it's a capitalised word skip checking because likely proper noun
                continue
            else:
                checked_word = spell.spell_checker(filetype, ilang, punc_free, gloss)
                if ilang == 'ru' and olang == 'en':
                    for key in gloss:
                        try:
                            pattern = str(key) + '*' #the pattern is * wildcard character + the key + * - so it matches wherever the key may be present in the string. Fnmatch is used so * is accepted as a wildcard
                            replacement = str(gloss[key]).lower()[2:-2] #the replacement string which contains the English translation of the Russian key - lowercased as the whole string will be
                            if fn.fnmatch(word, str(pattern)) == True: #fnmatch automatically casematches, so lowers un and the pattern - if there is a match, so if boolean = True
                                parag = parag.replace(punc_free, replacement)
                            elif fn.fnmatch(checked_word, str(pattern)) == True:
                                parag = parag.replace(punc_free, replacement)
                            else: #otherwise reset the loop and continue
                                parag = parag.replace(punc_free, checked_word)
                                
                        except Exception as e:
                            self.err_txt = str(e.args) + ', skipping word'
                            parag = parag.replace(punc_free, checked_word)
                            continue
                elif ilang == 'en' and olang =='zh_CN':
                    for key in gloss:
                        try:
                            pattern = str(key) + '*' #the pattern is * wildcard character + the key + * - so it matches wherever the key may be present in the string. Fnmatch is used so * is accepted as a wildcard
                            replacement = str(gloss[key]).lower()[2:-2] #the replacement string which contains the English translation of the Russian key - lowercased as the whole string will be
                            if fn.fnmatch(word, str(pattern)) == True: #fnmatch automatically casematches, so lowers un and the pattern - if there is a match, so if boolean = True
                                parag = parag.replace(punc_free, replacement)
                            elif fn.fnmatch(checked_word, str(pattern)) == True:
                                parag = parag.replace(punc_free, replacement)
                            else: #otherwise reset the loop and continue
                                parag = parag.replace(punc_free, checked_word)
                                
                        except Exception as e:
                            self.err_txt = str(e.args) + ', skipping word'
                            parag = parag.replace(punc_free, checked_word)
                            continue
                elif olang == 'ru' and ilang == 'en':
                    for key in gloss:
                        try:
                            pattern = '*' + str(gloss[key]) + '*'
                            replacement = str(key).lower()[2:-2]
                            if fn.fnmatch(word, str(pattern)) == True: #fnmatch automatically casematches, so lowers un and the pattern - if there is a match, so if boolean = True
                                parag = parag.replace(punc_free, replacement)
                            elif fn.fnmatch(checked_word, str(pattern)) == True:
                                parag = parag.replace(punc_free, replacement)
                            else: #otherwise reset the loop and continue
                                parag = parag.replace(punc_free, checked_word)
                                    
                        except Exception as e:
                            self.err_txt = str(e.args) + ', skipping word'
                            parag = parag.replace(punc_free, checked_word)
                            continue
                elif olang == 'en' and ilang == 'zh_CN':
                    for key in gloss:
                        try:
                            pattern = '*' + str(gloss[key]) + '*'
                            replacement = str(key).lower()[2:-2]
                            if fn.fnmatch(word, str(pattern)) == True: #fnmatch automatically casematches, so lowers un and the pattern - if there is a match, so if boolean = True
                                parag = parag.replace(punc_free, replacement)
                            elif fn.fnmatch(checked_word, str(pattern)) == True:
                                parag = parag.replace(punc_free, replacement)
                            else: #otherwise reset the loop and continue
                                parag = parag.replace(punc_free, checked_word)
                                    
                        except Exception as e:
                            self.err_txt = str(e.args) + ', skipping word'
                            parag = parag.replace(punc_free, checked_word)
                            continue
                else:
                    parag = parag.replace(punc_free, checked_word)
        
        return parag           
            
    def translate_data(self, ilanguage, olanguage, paragraph):
        try:
                transl_text = str(paragraph) #converts it into a string (to be sure that it is actually coming out as string value)
                transltxt = transl(source=ilanguage, target=olanguage).translate(transl_text) #translates the string - taken from deep_translator documentation
                #self.translated_doc.append(transltxt) #appends translated string to list in same order
                return transltxt
                #print(transltxt)
                
                 #so progress can be checked - and can see if it gets stuck on any particular entries
                #print(transltxt) #optional - can uncheck to see the translated results
                #print('translating2')
     
        except Exception as e: #except if there is an error 
            self.err_txt = str(e.args)
            
            
    def write_file(self, newdoc, location, filename, olang, filetype):
        print('location' + str(len(location)))
        langs_dict = transl.get_supported_languages(as_dict = True)
        lang_list = list(langs_dict.keys())
        abb_list = list(langs_dict.values())
        
        position = abb_list.index(olang)
        long_lang = lang_list[position]
        
        if newdoc == 'New document':
            ndoc = docx.Document()
            if len(str(location)) > 2:
                filen = os.path.basename(filename)
                nfile = filen[:-5] + '_translated_' + long_lang + '.docx'
                new_filename = os.path.join(location, nfile)
            else:
                if filetype == 'PDF':
                    
                    new_filename = str(filename)[:-4] + '_translated_' + long_lang + '.docx'
                else:
                    new_filename = str(filename)[:-5] + '_translated_' + long_lang + '.docx'
             
            
            file = r''+ new_filename
            self.ent_txt = 'New file created: ' + str(file)
            print(file)
        
            if len(self.translated_table) > 0:
                n=0
                for itable in self.translated_table:
                    rowlength = self.row_length[self.translated_table.index(itable)]
                    columnlength = self.column_length[self.translated_table.index(itable)]
                    df = pd.DataFrame(np.array(itable).reshape(rowlength, columnlength))
                    word_table = ndoc.add_table(rows = rowlength,cols = columnlength)
                    for x in range(0, rowlength):
                        for y in range(0, columnlength):
                            cell = word_table.cell(x,y)
                            try:
                                cell.text = df.iloc[x,y]
                            except:
                                cell.text = 'Error inputting'
                                continue
                            ndoc.save(file)
            else:
                pass
            
            if len(self.translated_doc) > 0:
                n=0
                for parag in self.translated_doc:
                    ndoc.add_paragraph(parag)
                    ndoc.save(file)
                    n = n + 1
                    self.chk3_txt = str(n) + ' paragraphs entered out of ' + str(len(self.translated_doc))
            else:
                pass
            
        else:
            n=0
            ndoc = docx.Document(filename)
            ndoc.add_paragraph('\n \n' + 'Translated to ' + str(long_lang.capitalize()) + '\n')
            ndoc.save(filename)
            self.ent_txt = 'Entering into existing file'
            
            
            if len(self.translated_table) > 0:
                n=0
                for itable in self.translated_table:
                    rowlength = self.row_length[self.translated_table.index(itable)]
                    columnlength = self.column_length[self.translated_table.index(itable)] 
                    df = pd.DataFrame(np.array(itable).reshape(rowlength, columnlength))
                    word_table = ndoc.add_table(rows = rowlength,cols = columnlength)
                    for x in range(0, rowlength):
                        for y in range(0, columnlength):
                            cell = word_table.cell(x,y)
                            try:
                                cell.text = df.iloc[x,y]
                            except:
                                cell.text = 'Error inputting'
                                continue
                            ndoc.save(filename)
            else:
                pass
            
            if len(self.translated_doc) > 0:
                n=0
                for parag in self.translated_doc:
                    ndoc.add_paragraph(parag)
                    ndoc.save(filename)
                    n = n + 1
                    self.chk3_txt = str(n) + ' paragraphs entered out of ' + str(len(self.untrans.paragraphs))
            else:
                pass
            
    def reset(self): #resets the function so there are no leftover values in the lists for the future values
        self.pdf_content.clear()
        self.tables.clear()
        self.translated_table.clear()
        self.row_length.clear()
        self.column_length.clear()
        self.doc.clear() #clears the untranslated list
        self.translated_doc.clear() #ibid for checked untranslated
        self.ldat_txt = '...'
        self.counter1_txt = '...'
        self.counter2_txt = '...'
        self.ent_txt = '...'
        self.chk3_txt = '...'
        
    
class TranslDocRun:
    def run(self, filename, filetype, location, ilanguage, olanguage, newdoc, windowx):
        
        TD = TranslateDoc()
        framea = tk.Frame(master=windowx, width = 150, height=100, padx=5, pady=5)
        framea.pack(fill = tk.X, side=tk.TOP)
        
        frameb = tk.Frame(master=windowx, width = 150, height = 100, padx=5, pady=5)
        frameb.pack(fill=tk.X, side = tk.TOP)
        
        framec = tk.Frame(master=windowx, width = 150, height = 100, padx=5, pady=5)
        framec.pack(fill=tk.X, side = tk.TOP) 
        
        fl = tk.Label(master=framea, text = 'Now translating:' + str(filename))
        fl.pack(fill = tk.X, side = tk.TOP)
        
        err = tk.Label(master=framea, text = '...')
        err.pack(fill=tk.X, side = tk.TOP)
        
        counter0 = tk.Label(master=framea, text = '...')
        counter0.pack(fill = tk.X, side = tk.BOTTOM)
        
        ldat = tk.Label(master=framea, text = '...')
        ldat.pack(fill = tk.X, side = tk.BOTTOM)
    
        counter1 = tk.Label(master=frameb, text = '...')
        counter1.pack(fill=tk.X, side = tk.BOTTOM)
        
        framed = tk.Frame(master=windowx, width = 150, height=100, padx=5, pady=5)
        framed.pack(fill = tk.X, side=tk.TOP)
        
        framee = tk.Frame(master=windowx, width = 150, height=100, padx=5, pady=5)
        framee.pack(fill = tk.X, side=tk.TOP)
        
             
        counter2 = tk.Label(master=framec, text = '...')
        counter2.pack(fill=tk.X, side = tk.BOTTOM)
    
        ent = tk.Label(master=framed, text = '...')
        ent.pack(fill = tk.X, side = tk.TOP)
        
        chk3 = tk.Label(master=framee, text = '...')
        chk3.pack(fill = tk.X, side = tk.BOTTOM)
        
        
        windowx.after(1000, TD.update, err, counter0, ldat, counter1, counter2, ent, chk3, windowx)
        
        TD.load_data(filename, filetype, ilanguage)
        TD.check_data(filetype, ilanguage, olanguage)
        TD.write_file(newdoc, location, filename, olanguage, filetype)
        framea.forget()
        frameb.forget()
        framec.forget()
        framed.forget()
        framee.forget()
        TD.reset()