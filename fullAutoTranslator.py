'''
Created on 23 Jun 2022

@author: turch
'''

import tkinter as tk
from tkinter import *
from deep_translator import GoogleTranslator as transl #uses deep_translator instead of googletrans because
# googletrans maxes out times that you can run entries - this has a much higher threshold
#although if it starts timing out lots can change to MicrosoftTranslator - still effective 
import os
import pandas as pd
import traceback
import threading
import fnmatch as fn
import docx
import re
import numpy as np
import pytesseract
import pdf2image
from pdf2image import convert_from_path
import PIL
import xlwings as xw
import enchant.checker
import difflib
import spellchecker as sc
import PyPDF2
'''
United into one module so it can be easily converted into an .exe application.
CONTENTS:
line 79 - class GUI
    line 89 - start
    line 107 - runrun
    line 115 - path_window
line 189 - class loadFolders
    line 196 - openfolder
    line 301 - makesql (crossed out)
    line 334 - runtransl
    line 444 - removefiles
    line 476 - input_window
    line 527 - output_location
    line 564 - runtransl2
    line 589 - xl_proceed
    line 684 - doc_proceed
line 717 - class Run
    run - line 718
class TranslMethods - line 729
    update - line 749
    load_data - line 762
    data_check - line 791
    translator - line 922
    combinedata - line 948
    input_data - line 975
    sql_work - line 1005 (crossed out)
    reset - line 1030
class TranslateRun - line 1046
    runn - line 1049
    runtransl - line 1061
class TranslDocMethods - line 1128
    update - line 1145
    load_data - line 1155
    check_data - line 1242
    check_data2 - line 1309
    translate - line 1391
    write_file - line 1407
    reset - line 1500
class TranslDocRun - line 1515
    run - line 1516
class Spelling - line 1572
    ru_dict - line 1575
    create_glossary - line 1600
    xl_spellcheck - line 1632
    doc_spellcheck - line 1693
    spellchecker - line 1756
class ResetGUI - line 1871
    restart - line 1872
    reset - line 1901
'''

class GUI:
    '''
    The goal here is to create a GUI for the Translator program with user input. First the start window is created welcoming the user and prompting them to click Start. The data_path window then appears and it is input. Once 
    the data path is put in, the openfolder function can be called from run
    '''
    paths = 0 #creates path variable
    input_language = 'ru'
    output_language = 'en'
    filetype = 'Excel'
    
    def start(self):
        global window
        window = tk.Tk()
        frame = tk.Frame(master=window, width = 150, height=150)
        frame.pack() #makes a frame and packs it to put widgets within
        
        def enter_path(): #closes existing window and opens the set path
            frame.destroy()
            GUI.path_window(self)
            
        
        
        greeting = tk.Label(master=frame, text="Welcome to the CAAL Auto-translator for all your auto-translation file needs. \n \n Make sure needed files are closed before starting, the program will not work if they are open")
        start_btn = tk.Button(master=frame, text = 'START', command=enter_path) #
        greeting.pack(fill = tk.X)
        start_btn.pack(fill = tk.X)
        window.mainloop()
    
    def runrun(self):
        lf = Run()
        pinput = GUI.paths #takes the first entry of paths
        window.destroy()
        #return str(pinput), str(GUI.input_language), str(GUI.output_language), str(GUI.filetype)
        lf.run(pinput, str(GUI.input_language), str(GUI.output_language), str(GUI.filetype))
        
        
    def path_window(self): #codes a window for entering data path, setting language of data and to translate into
        frame = tk.Frame(master=window, width = 300, height=300)
        frame.pack(fill = tk.X, side = tk.TOP)
        
        path = tk.Label(master=frame, text = 'Enter data path:')
        path_input = tk.Entry(master=frame)
        path.pack(fill = tk.X, side=tk.LEFT)
        path_input.pack(fill = tk.X, side = tk.RIGHT)
        
        framea = tk.Frame(master=window, width=150, height=150)
        framea.pack(fill = tk.X, side =tk.TOP)
        
        langs_dict = transl.get_supported_languages(as_dict = True)
        lang_list = list(langs_dict.keys())
        abb_list = list(langs_dict.values())
        
        #print(langs_dict)
        
        def input_lang(selection):
            position = lang_list.index(str(ilang.get()).lower())
            GUI.input_language = abb_list[position]
            #print(GUI.input_language)
            
            
        def output_lang(selection):
            position = lang_list.index(str(olang.get()).lower())
            GUI.output_language = abb_list[position]
            #print(GUI.output_language)
            
        def fltype(selection):
            GUI.filetype = ftype.get()
        
        ilang = StringVar(framea)
        ilang.set("Russian")
        ilanguage = OptionMenu(framea, ilang, 'Russian', 'English', 'Chinese (Simplified)', 'Kazakh','Kyrgyz', 'Tajik', 'Turkmen', 'Uzbek', 'Italian', 'French', 'Spanish','German', 'Arabic', 'Ukrainian', command = input_lang)
        lang = tk.Label(master=framea, text = 'Enter language of data, set to Russian by default:')
        lang.pack(fill = tk.X, side = tk.LEFT)
        ilanguage.pack(fill = tk.X, side = tk.RIGHT)
        
        frameb = tk.Frame(master=window, width=150, height = 150)
        frameb.pack(fill = tk.X, side = tk.TOP)
        
        lang2 = tk.Label(master = frameb, text = 'Enter language to be translated into, English as default:')
        lang2.pack(fill = tk.X, side = tk.LEFT)
        olang = StringVar(frameb)
        olang.set("English")
        olanguage = OptionMenu(frameb, olang, 'Russian', 'English', 'Chinese (Simplified)', 'Kazakh','Kyrgyz', 'Tajik', 'Turkmen', 'Uzbek', 'Italian', 'French', 'Spanish','German', 'Arabic', 'Ukrainian', command = output_lang)
        olanguage.pack(fill = tk.X, side = tk.RIGHT)
        
        framec = tk.Frame(master=window, width =150, height = 150)
        framec.pack(fill = tk.X, side = tk.TOP)
        
        typefile = tk.Label(master=framec, text = 'Select type of file to be translated, Excel as default:')
        typefile.pack(fill = tk.X, side = tk.LEFT)
        ftype = StringVar(framec)
        ftype.set('Excel')
        filetypes = OptionMenu(framec, ftype, 'Excel', 'PDF', 'Word', command = fltype)
        filetypes.pack(fill = tk.X, side = tk.RIGHT)
        
        framed = tk.Frame(master = window, width = 150, height = 150)
        framed.pack(fill = tk.X, side = tk.TOP)
        
        note = tk.Label(master=framed, text ='NOTE: Currently only one type of file can be translated into one language \n at a time, in future multiple types will be possible, \n for now just start a new run for different filetypes/languages')
        note.pack(fill=tk.X, side =tk.TOP)
        
       
        def next_window(): #defines next_window function
            GUI.paths = str(path_input.get())
            GUI.runrun(self) #runs
            
            
        ok_btn = tk.Button(master=framed, text = 'OK', command=next_window)
        ok_btn.pack(fill = tk.X, side = tk.BOTTOM)
        
        
        
        window.mainloop()


class loadFolders: #creates a class to load folders within the given path for translation of all files within that folder
    archive = [] #creates list for archive spreadsheets
    monument = [] #and list for monument ones
    other = []
    docs = []
    done = []
    biglist = [archive, monument, other] #biglist for the excel spreadsheet lists so they can be cycled through
    l = None
    
    def openfolder(self, data_path, filetype): 
        '''
        This function is the first to be executed in the program as it locates the folders from the specified data path, accesses the files within them, and extracts the file paths to the specified lists (optional ultimately
        but useful for this group of spreadsheets). The file paths can then be used to access individual files for translation.
        '''
        self.monument.clear()
        self.archive.clear()
        self.docs.clear()
        
        if '"' in str(data_path):
            data_path = str(data_path).replace('"', '')
            #print('w:)')
            
        else:
            #print('w:(')
            pass
            
        if filetype == 'Excel':
                if 'archive' in str(data_path).lower():
                    #print('i')
                    if '.xl' in str(data_path):
                        self.archive.append(str(data_path))
                    for data_path, subdirs, files in os.walk(data_path):
                        for f in files:
                            file_path = os.path.join(data_path, f)
                            self.archive.append(str(file_path))
                        for s in subdirs:
                            subdir_path = os.path.join(data_path, s)
                            #print(subdir_path)
                            for f in files: #for the directory, subdirectories, and files in the directory path
                                file_path = os.path.join(subdir_path, f) #join file name with the directory path as file path
                                #print(file_path)
                                self.archive.append(str(file_path)) #and append that to the archive list 
            
                elif 'monument' in str(data_path).lower():
                    #print('x')
                    if '.xl' in str(data_path):
                        self.monument.append(str(data_path))
                    else:
                        for data_path, subdirs, files in os.walk(data_path):
                            for f in files:
                                file_path = os.path.join(data_path, f)
                                self.monument.append(str(file_path))
                            for s in subdirs:
                                subdir_path = os.path.join(data_path, s)
                                #print(subdir_path)
                                for f in files: #for the directory, subdirectories, and files in the directory path
                                    file_path = os.path.join(subdir_path, f) #join file name with the directory path as file path
                                    #print(file_path)
                                    self.monument.append(str(file_path)) #and append that to the archive list 
                        
                        
                else:
                    if '.xl' in str(data_path):
                        #print('yi')
                        self.other.append(data_path)
                    else:
                        for data_path, subdirs, files in os.walk(data_path):
                            for f in files:
                                file_path = os.path.join(data_path, f)
                                self.other.append(f)
                            for s in subdirs: #for each subdirectory
                                dirpath = os.path.join(data_path, s) #creates a directory path to add the file name to for a complete file path
                                if 'archive' in str(dirpath).lower(): #if 'ARCHIVE' is in the subdirectory name
                                    for dirpath, subdirs, files in os.walk(dirpath): #for the directory, subdirectories, and files in the directory path
                                        for f in files:
                                            file_path = os.path.join(dirpath, f)
                                            if str(file_path) in self.archive:
                                                continue
                                            else:
                                                self.archive.append(file_path)
                                            
                                            
                                elif 'monument' in str(dirpath).lower(): #otherwise if 'MONUMENT' is in the subdirectory path:
                                    for dirpath, subdirs, files in os.walk(dirpath): #do the exact same process but append it to the monument list instead
                                        for f in files:
                                            file_path = os.path.join(dirpath, f)
                                            if str(file_path) in self.monument:
                                                continue
                                            else:
                                                self.monument.append(file_path)
                                            
                                else: #otherwise continue
                                    for data_path, subdirs, files in os.walk(data_path):
                                        for f in files:
                                            file_path = os.path.join(data_path, f)
                                            self.other.append(file_path)
                                        for s in subdirs:
                                            dirpath = os.path.join(data_path, s)
                                            for dirpath, subdirs, files in os.walk(dirpath):
                                                for f in files:
                                                    file_path = os.path.join(dirpath, f)
                                                    if str(file_path) in self.other:
                                                        continue
                                                    else:
                                                        self.other.append(file_path)
                                            
        else:
            if '.docx' in str(data_path) or '.pdf' in str(data_path):
                self.docs.append(data_path)
                dirpath = os.path.dirname(data_path)
                self.dirpath = dirpath
            elif '.xl' in str(data_path):
                if 'archive' in str(data_path).lower():
                    self.archive.append(data_path)
                elif 'monument' in str(data_path).lower():
                    self.monument.append(data_path)
                else:
                    self.other.append(data_path)
            else:
                for data_path, subdirs, files in os.walk(data_path):
                    for f in files:
                        file_path = os.path.join(data_path, f)
                        if str(file_path) in self.docs or '.xls' in str(file_path):
                            continue
                        else:
                            self.docs.append(str(file_path))
                    if len(list(subdirs)) >0:
                        for s in subdirs:
                            dirpath = os.path.join(data_path, s)
                            self.dirpath = dirpath
                            for dirpath, subdirs, files in os.walk(dirpath):
                                for s in subdirs:
                                    dirpath2 = os.path.join(data_path, s)
                                    for dirpath, files in os.walk(dirpath2):
                                        for f in files:
                                            file_path = os.path.join(dirpath2, f)
                                            if str(file_path) in self.docs or '.xls' in str(file_path):
                                                continue
                                            else:
                                                self.docs.append(str(file_path))
                                for file in files:
                                    file_path = os.path.join(dirpath, file)
                                    if str(file_path) in self.docs or '.xls' in str(file_path):
                                        continue
                                    else:
                                        self.docs.append(str(file_path))
                    else:
                        for file in files:
                            self.dirpath = data_path
                            file_path = os.path.join(data_path, file)
                            if str(file_path) in self.docs or '.xls' in str(file_path):
                                continue
                            else:
                                self.docs.append(str(file_path))
                
          
    ''' 
    def makesql(self, data_path): #makes an SQL from sqlanalysis module'''
    '''
        This is the first step to loading all the data into an SQL for further analysis. Because we are iterating through files, but we want to compare across files, the SQL database is made now
        so that the tables already exist for the folders, but will be filled with data for each file in the translate methods
        ''''''
        folder = 'CAAL SQL' #the database name
        self.folder = folder
        
        ascii_conv = [] #creates list for ascii converted values
        for c in str(data_path): #for character in data_path
            ascii_conv.append(ord(c)) #find its ascII value and append it to the table
        uniq_table = sum(ascii_conv) #uniq_table value is the sum of all the ascII
        
        
        tables = [str('Archive' + str(uniq_table)), str('Monument' + str(uniq_table))] 
        
        ''''''the tables for archive and monument data - the uniq_table value is added to make each table name unique so that
        they don't have to be deleted across different running of the program unless they are the same. I wanted to find a value that would be unique
        for each data path across different times the program is run without having to use the data_path itself as that is quite long. So by getting all the  
        ASCII values and adding them this will be a unique value because the data paths will vary in at least one character. This will allow to save the
        data obtained from the translated spreadsheets within an SQL database, with the files within one folder stored within one table.
        Then when more files have been translated the tables can all be united in a separate program and extensive data analysis can be run
        For now the data analysis is very rudimentary as the file translation is still being optimised''''''
        self.tables = tables
        
        SQLcolumns = 'Row TEXT, Name TEXT, Description TEXT, Location TEXT' #name, description, and location files - will be used to run some analysis of how many entries come from where, what type they are (based on common words)
        self.SQLcolumns = SQLcolumns
        
        lite = SQLite(folder) #lite is the class SQLite in sqlanalysis - runs the init module with CAAL SQL as the database
        for table in tables: #for each string in tables
            lite.makesql(table, SQLcolumns) #make a table with that as the name, and the given columns
        
         '''       
    def runtransl(self, ilanguage, olanguage, data_path, window, filetype): #runs the translation methods defined in the translate module
        spell = Spelling() #spell is the Spelling() class from spellcheck
        spell.ru_dict() #runs the making russian dictionary module from spelling()

        
        
        framea = tk.Frame(master=window, width = 300, height=300)
        framea.pack(fill = tk.X, side = tk.TOP)
        
        def alterfile():
            global remove #so remove can be called elsewhere
            remove = remo.get()
            framea.destroy()
            frameb.destroy()
            self.removefiles(remove, l, data_path, ilanguage, olanguage, window, filetype) #calls removefiles function
        
        def resett():
            reset = ResetGUI()
            window.destroy()
            reset.restart()
        
        if len(self.archive) == 0 and len(self.monument) == 0 and len(self.docs) == 0 and len(self.other) == 0:
            empty = tk.Label(master=framea, text = 'Empty input! Did you select the wrong filetype or put in an invalid path? Reset and try again.')
            empty.pack(fill = tk.X, side = tk.TOP)
            
            reset_btn = tk.Button(master=framea, text = 'RESET', command = resett)
            reset_btn.pack(fill = tk.X, side = tk.BOTTOM)
            
        def list_files(l, textt):
            self.done.append(l)
            self.l = l 
            arc = tk.Label(master=framea, text = textt) #create label widget saying that and listing files
            arc.pack(fill = tk.X, side = tk.TOP)
            
            
            for f in l: #for each entry in l
                n = int(l.index(f)) + 1
                name = str(str(n) + '. ' + str(f)) #file name is the path + number beforehand
                lst = tk.Label(master=framea, text = name) #creates label so files are listed in window
                lst.pack(fill = tk.X, side = tk.TOP)
        
        if filetype == 'Excel':
                
            for l in self.biglist: #for l in biglist
                if len(l) == 0 and len(self.done) == 2:
                    GUI = ResetGUI()
                    GUI.restart()
                elif len(l) == 0 and len(self.done) < 2:
                    self.done.append(l)
                    continue
                else:
                    if l == self.archive: #if l is archive
                        textt = 'The folder being translated is Archives. The files in this folder are:'
                        list_files(l, textt)
                    elif l == self.monument:
                        textt = 'The folder being translated is Monuments. The files in this folder are:'
                        list_files(l, textt)
                    else:
                        textt = 'The files in this folder are:'
                        list_files(l, textt)
                        
                    #remove = input("\n If you want to remove any spreadsheets, enter the corresponding numbers here separated by comma, otherwise enter N:") 
                    frameb = tk.Frame(master=window, width = 300, height=150)
                    frameb.pack(fill = tk.X, side = tk.TOP)
                    remy = tk.Label(master=frameb, text = 'If you want to remove any spreadsheets, enter corresponding numbers here separated by comma, otherwise enter N')
                    remy.pack(fill = tk.X, side = tk.TOP) #creates label for entry to begin file removal process if neede
                    global remo #makes remo global so it can be called in other functions
                    remo = tk.Entry(master=frameb) #remo is the entry
                    remo.pack(fill = tk.X, side = tk.TOP) #packs it
                    
                    ok_btn = tk.Button(master=frameb, text = 'OK', command=alterfile)
                    ok_btn.pack(fill = tk.X, side = tk.BOTTOM)
                    window.mainloop()
                        
        elif filetype == 'Word' or filetype == 'PDF':
            l = self.docs
            textt = 'The files to be translated are:'
            list_files(l, textt)
            
            frameb = tk.Frame(master=window, width = 300, height=150)
            frameb.pack(fill = tk.X, side = tk.TOP)
            remy = tk.Label(master=frameb, text = 'If you want to remove files, enter corresponding numbers here separated by comma, otherwise enter N')
            remy.pack(fill = tk.X, side = tk.TOP) #creates label for entry to begin file removal process if neede
            remo = tk.Entry(master=frameb) #remo is the entry
            remo.pack(fill = tk.X, side = tk.TOP) #packs it
            
            ok_btn = tk.Button(master=frameb, text = 'OK', command=alterfile)
            ok_btn.pack(fill = tk.X, side = tk.BOTTOM)
            window.mainloop()
            
        else:
            self.biglist.append(self.docs)
            for l in self.biglist:
                if len(l) == 0 and len(self.done) == 3:
                    GUI = ResetGUI()
                    GUI.restart()
                elif len(l) == 0 and len(self.done) < 3:
                    self.done.append(l)
                    continue
                else:
                    if l == self.archive: #if l is archive
                        textt = 'The folder being translated is Archives. The files in this folder are:'
                        list_files(l, textt)
                    elif l == self.monument:
                        textt = 'The folder being translated is Monuments. The files in this folder are:'
                        list_files(l, textt)
                    else:
                        textt = 'The files in this folder are:'
                        list_files(l, textt)
                        
                frameb = tk.Frame(master=window, width = 300, height=150)
                frameb.pack(fill = tk.X, side = tk.TOP)
                remy = tk.Label(master=frameb, text = 'If you want to remove files, enter corresponding numbers here separated by comma, otherwise enter N')
                remy.pack(fill = tk.X, side = tk.TOP) #creates label for entry to begin file removal process if neede
                remo = tk.Entry(master=frameb) #remo is the entry
                remo.pack(fill = tk.X, side = tk.TOP) #packs it
                
                ok_btn = tk.Button(master=frameb, text = 'OK', command=alterfile)
                ok_btn.pack(fill = tk.X, side = tk.BOTTOM)
                window.mainloop()
                    
                
    def removefiles(self, remove, l, data_path, ilanguage, olanguage, window, filetype):
        '''
        This allows the user to remove any files they want by listing their index - this is useful for files that have already been translated, or that you know have problems, or want to avoid for whatever reason
        The program currently successfully iterates through at least 3 - 4 files. If the program ends up being slow, it might be worth running the shorter-entry files to showcase that it is working
        '''
        #testff = tk.Label(master=framea, text = str(remove) + 'works')
        #testff.pack(fill= tk.X, side = tk.TOP)
        if str(remove) == 'N': #if No files are to be removed
            if filetype == 'Excel':
                self.input_window(ilanguage, olanguage, window, filetype, data_path) #move to input_window
            elif filetype == 'Word' or 'PDF':
                self.output_location(filetype, ilanguage,olanguage,window)
            else:
                if self.l == self.archive or self.l == self.monument or self.l == self.other:
                    self.input_window(ilanguage, olanguage, window, filetype, data_path)
                else:
                    self.output_location(filetype, ilanguage, olanguage, window)
              
        elif str(remove) == '': #if it's a blank (clicked OK by accident)
            self.runtransl(ilanguage, olanguage, data_path, window, filetype) #rerun the window
        else: #otherwise
            if ',' in str(remove): #if , is in the string
                rem = remove.split(',') #split values by presence of comma
                i = 1 #i is 0
                for r in rem: #for r in the split string
                    del l[int(r) - int(i)] #delete the file from list using its index calculated by subtracting i from the number given (because the index will update due to deletion i needs to be updated as well)
                    i += 1 #because every time a file is deleted the remaining indices are updated i has to be updated as well
                self.runtransl(ilanguage, olanguage, data_path, window, filetype) #runs window with updated files
            elif str(remove).isdigit() == True:
                del l[int(str(remove)) - 1]
                self.runtransl(ilanguage, olanguage, data_path, window, filetype)
            else:
                windowc = tk.Tk()
                inv = tk.Label(windowc, text = 'Invalid Input')
                inv.pack(fill = tk.X, side = tk.BOTTOM)
                self.runtransl(ilanguage, olanguage, data_path, window, filetype)
                windowc.mainloop()
    
    def input_window(self, ilanguage, olanguage, windowa, filetype, data_path):
        framea = tk.Frame(master=windowa, width = 150, height=150, padx=5, pady=5)
        framea.pack(fill = tk.X, side=tk.TOP)
        
        sheet = tk.Label(master=framea, text = 'Please enter sheet:')
        sheet_input = tk.Entry(master=framea, width = 75)
        sheet.pack(fill = tk.X, side=tk.LEFT)
        sheet_input.pack(fill = tk.X, side = tk.RIGHT)
        
        frameb = tk.Frame(master=windowa, width = 150, height=150, padx=5, pady=5)
        frameb.pack(fill = tk.X, side=tk.TOP)
        
        col1 = tk.Label(master=frameb, text = 'Please enter columns to be translated with slash if you want them to be translated as one, and with a comma if you want them to be translated separately \n (ex: H/J will translate columns H and J and combine them; AK, A will translate columns AK and A separately)')
        col1_input = tk.Entry(master=frameb, width = 75)
        col1.pack(fill = tk.X, side=tk.LEFT)
        col1_input.pack(fill = tk.X, side = tk.RIGHT)
        
        framec = tk.Frame(master=windowa, width = 150, height=150, padx=5, pady=5)
        framec.pack(fill = tk.X, side=tk.TOP)
        
        col2 = tk.Label(master=framec, text = 'Enter input columns in same order/format that you did columns (if you want data from H/J to be input into I, put I first): \n Note: if same column is put in for input/output, it will enter the translated data into the same cell as untranslated and will keep both')
        col2_input = tk.Entry(master=framec, width = 75)
        col2.pack(fill = tk.X, side=tk.LEFT)
        col2_input.pack(fill = tk.X, side = tk.RIGHT)
        
        framed = tk.Frame(master=windowa, width = 150, height=150, padx=5, pady=5)
        framed.pack(fill = tk.X, side=tk.TOP)
        
        col3 = tk.Label(master=framed, text = 'Enter row number from which to start translation (ex: 5 - exclude column names):')
        col3_input = tk.Entry(master=framed, width = 75)
        col3.pack(fill = tk.X, side=tk.LEFT)
        col3_input.pack(fill = tk.X, side = tk.RIGHT)
        
        framee=tk.Frame(master=windowa, width=150, height=150, padx=5, pady=5)
        framee.pack(fill=tk.X, side = tk.TOP)
        
        frame_list = [framea,frameb,framec,framed,framee]
        
        def get_inputs():
            input_sheet = sheet_input.get()
            input_columns = col1_input.get()
            output_columns = col2_input.get()
            start_row = col3_input.get()
            for f in frame_list:
                f.forget()
            self.runtransl2(windowa, ilanguage, olanguage, input_sheet, input_columns, output_columns, start_row, filetype, data_path)
            
        
        ok_btn = tk.Button(master=framee,text = 'OK',command=threading.Thread(target=get_inputs).start)
        ok_btn.pack(fill = tk.X, side = tk.BOTTOM)
            
    def output_location(self, filetype, ilanguage, olanguage, window):
        framea = tk.Frame(master=window, width = 150, height=150, padx=5, pady=5)
        framea.pack(fill = tk.X, side=tk.TOP)
        
        def ndd(selection):
            global newdoc
            newdoc = nd.get()
        
        new_doc = tk.Label(master=framea, text = 'Please enter if translation is in new document or within same one (Word only): \n Note: new documents are output as same filename with added "_translated" and new language')
        new_doc.pack(fill = tk.X, side = tk.LEFT)
        nd = StringVar(framea)
        nd.set("-")
        ndmenu = OptionMenu(framea, nd, 'New document', 'Write in same one', command = ndd)
        ndmenu.pack(fill = tk.X, side = tk.RIGHT)
        
        frameb = tk.Frame(master=window, width = 150, height=150, padx=5, pady=5)
        frameb.pack(fill = tk.X, side=tk.TOP)
        framec = tk.Frame(master=window, width = 150, height = 50)
        framec.pack(fill = tk.X, side = tk.TOP) 
        
        def locationn():
            location = loc.get()
            #print(location)
            framea.forget()
            frameb.forget()
            framec.forget()
            transl = TranslDocRun()
            if len(str(location)) == 0:
                locat = self.dirpath
            else:
                locat = str(location)
            for file in self.docs:
                window.update()
                transl.run(window, file, filetype, locat, ilanguage, olanguage, newdoc) #frameb,framec,framed,framee,err, ldat, counter0,counter1,counter2,ent,chk3)
            self.docs.clear()  
            reset = ResetGUI()
            reset.restart()
            
        locate = tk.Label(master=frameb, text = 'Enter folder to which new file is saved, leave blank if same as original file:')
        locate.pack(fill = tk.X, side = tk.LEFT)
        loc = tk.Entry(master=frameb, width = 75)
        loc.pack(fill = tk.X, side = tk.RIGHT)
           
        ok_btn = tk.Button(master=framec, text = 'OK', command=threading.Thread(target=locationn).start)
        ok_btn.pack(fill = tk.X, side = tk.BOTTOM)
                
    def runtransl2(self, windowa, ilanguage, olanguage, input_sheet, input_columns, output_columns, start_row, filetype, data_path):
        
        columns = input_columns
        col_list = columns.split(',') #creates list of columns with , as the splitter

        input_col = output_columns
        inputs = input_col.split(',') #separates input columns based on ,
    
        columndict = dict(zip(col_list,inputs)) #zips the list of columns and inputs together so the 1st column/group of columns becomes associated with the 1st input (as key/value) and so on
        
        self.xl_proceed(windowa, columndict, input_sheet, start_row, ilanguage, olanguage, col_list, inputs, filetype, data_path)
        
        '''
        
            Here the user is able to put in which sheet the program should access and which columns they would like to translate and in what order. This is left up to the user so that if only a particular 
            column needs to be translated quickly for multiple files, that can be accounted for as well as, having all the needed fields be translated by putting in multiple columns. The format and order are critical here:
            the format because there are columns like AK or BD made up of 2 characters that still need to be read as one entry, but there is also a need to combine data from 2 Russian-entry columns into 1 English-entry
            column. As such if a / is placed then the columns are translated together and output as 1 entry; if a comma is placed, then that means the columns should be translated separately with separate outputs. T
            The order has to be the same for the input and output as that is the only way the program knows which column corresponds to which
            
        '''  
    # table = self.tables[0] #once each file is cycled through the SQL table can be accessed for analysis - in this case the Archive table is used but either can be depending on what
        #is wanted for the analysis
    # SQLite.analysis(SQLite, table, self.SQLcolumns) #call SQLite analysis to find number of different monuments
    
    def xl_proceed(self, windowa, column_dict, input_sheet, start_row, ilanguage, olanguage, col_list, inputs, filetype, data_path):
        transl = TranslateRun()
        
        windowb = tk.Tk()
        framea = tk.Frame(master=windowb, width = 150, height=150, padx=5, pady=5)
        framea.pack(fill = tk.X, side=tk.TOP)
        
        txt1 = 'Sheet: ' + str(input_sheet) + '\n'
        
        for col in col_list:
            c = int(col_list.index(col))
            txt2 = 'Column to translate: ' + str(col) + '   Column for input: ' + str(inputs[c]) + '\n'
            txt1 = txt1 + txt2
         
        txt = txt1 + '\n Starting row: ' + str(start_row) + '\n Confirm CONTINUE or go BACK:'
        confrm = tk.Label(master=framea, text = str(txt))
        confrm.pack(fill = tk.X, side = tk.TOP)
        
        frameb = tk.Frame(master=windowb, width =150, height=75, padx=5, pady=5)
        frameb.pack(fill = tk.X, side=tk.TOP)
        
        def go_back():
            windowb.destroy()
            self.input_window(ilanguage, olanguage, windowa, filetype, data_path)
            
        def runn():
            windowb.destroy()
            
            try:     
                for file in self.l:
                    
                    #print('worrk')
                                     
                    if file in self.archive:
                    #setting which row to start on - this can easily also be a user input
                        cols = 'D' #setting which column is used for counter - it has to be a column which will always have data in it regardless of entry so it is set to the CAAL ID column which has 
                    #to be filled for the entry to exist in the spreadsheet (as if it has no CAAL ID it should not be getting recorded in the first place) 
                    elif file in self.monument:
                    #ibid - setting it here but can easily be changed through user input
                        cols = 'G' #sets G as the column because it is the CAAL ID
                    
                    else:
                        cols = 'A'
                    row_length = pd.read_excel(file, input_sheet, usecols = cols) #creates panda dataframe to figure out how many rows there are in the sheet
                    #column D is used bc for both types of spreadsheets it is a pre-filled column that has to be filled for the row to exist
                    #which is then needed to find the max_row so it can be used in ranges
                    max_row = int(len(row_length.dropna())) + 3#defines max_row (to be used in ranges as the length of row_length dataframe - as the length of the dataframe with the extracted column)
                    #dropna is needed because many of the spreadsheets have hundreds of thousands of blank cells loaded in below the data that end up being counted otherwise which makes everything MUCH slower
                    #but because it drops any blank values the column for it has to be one that will have values for every single entry - so the CAAL_ID is chosen an entry has to have a CAAL ID to be entered
                    #in the first place
                    
                    for key in column_dict: #for each key in the column dictionary (where the data columns and input columns are
                        column_names = []
                        input_column = str(column_dict[key])
                        if '/' in key:
                            column_names = key.split('/')
                            
                
                        else:
                                column_names.append(key)
                        
                        ##print('working')
                        transl.runtransl(file, input_sheet, column_names, input_column, int(start_row), int(max_row), ilanguage, olanguage, windowa, filetype)
                    #transl.SQLinput(file, self.folder, input_sheet, self.tables, data_cols, max_row)
                reset = ResetGUI()
                
                if len(self.done) == 3:
                    self.biglist.clear()
                    self.done.clear()
                    #print('all 3 added')
                    reset.restart()
                else:
                    self.biglist.remove(self.l)
                    windowc = tk.Tk()
                    self.runtransl(ilanguage, olanguage, data_path, windowc, filetype)
                        
                    
            except Exception as e:
                    reset = ResetGUI()
                    windoww = tk.Tk()
                    fram = tk.Frame(master=windoww, width = 150, height = 150, padx = 5, pady=5)
                    fram.pack(fill = tk.X, side = tk.TOP)
                    err = tk.Label(master=fram, text = 'ERROR:' + str(e) + ' ' + str(e.args) + '/n' + str(traceback.print_exc(limit=None, file=None, chain=True)))
                    err.pack(fill=tk.X, side = tk.TOP)
                    tip = tk.Label(master=fram, text = 'If Permission Error check permissions on files and tick "allowed to edit" for everyone or check that files are not open')
                    tip.pack(fill=tk.X, side = tk.TOP)
                    
                    def resett():
                        windoww.destroy()
                        reset.restart()
                        
                    framee = tk.Frame(master=windoww, width = 150, height = 150, padx = 5, pady=5)
                    framee.pack()
                    new_btn = tk.Button(master=framee, text = 'RESTART', command = resett)
                    new_btn.pack(fill = tk.X, side = tk.BOTTOM)
                    windoww.mainloop()
            
        
        back = tk.Button(master=frameb, text = 'BACK', command=go_back)
        back.pack(fill = tk.Y, side = tk.LEFT)
        cont = tk.Button(master =frameb, text = 'CONTINUE', command=runn)
        cont.pack(fill = tk.Y, side = tk.RIGHT)
        windowb.mainloop()
    
    def doc_proceed(self, filetype, ilanguage, olanguage, newdoc, location, window):
        transl = TranslDocRun()
        reset = ResetGUI()
        framea = tk.Frame(master=window, )
        
        def runn():
            windowz = tk.Tk()
            for file in self.docs:
                txt = 'Now translating: ' + str(file)
                
                transl.run(windowz, file, filetype, location, ilanguage, olanguage, newdoc, txt) #frameb,framec,framed,framee,err, ldat, counter0,counter1,counter2,ent,chk3)
                
            self.docs.clear()  
            reset.restart()
            windowz.mainloop() 
        
        ok_btn = tk.Button(master=framea, text = 'CONFIRM', command=runn())
        ok_btn.pack(fill=tk.X, side =tk.BOTTOM)
            
        
class Run: 
    def run(self, data_path, ilanguage, olanguage, filetype):
            window = tk.Tk()
            #data_path = data #str(input("Enter full main folder path here: ")) - this is normally a user input but I've made it a preset path for assignment purposes
            #language = 'ru' #can change this to Turkmen/Uzbek/Chinese - in future will be dropdown list for user input
            
            load = loadFolders() #the load folder class
            load.openfolder(data_path, filetype) #opens the folders in the data path
            #load.makesql(data_path)
            load.runtransl(ilanguage, olanguage, data_path, window, filetype) #runs translation
            window.mainloop()

class TranslMethods:
    
    #untranslated = [] #creates list for untranslated strings
    checked_untrans = [] #creates list for untranslated strings checked against the glossary
    translated = [] #creates list for translated strings
    combineddata = []
    rows = 0
    err_txt = 'Error status: Functioning fine'
    ldat_txt = '...'
    extr_txt = '...'
    chk_txt = '...'
    counter1_txt = '...'
    chk2_txt = '...'
    counter2_txt = '...'
    com_txt = '...'
    ent_txt = '...'
    chk3_txt = '...'
    resett = False
    txt_lst = [err_txt, ldat_txt, extr_txt, chk_txt, counter1_txt, chk2_txt, counter2_txt, com_txt, ent_txt, chk3_txt]
    
    def update(self, err, ldat, extr, chk, counter1, chk2, counter2, com, ent, chk3, windowx): #function to allow window to continue updating
        err.config(text = self.err_txt)
        ldat.config(text = self.ldat_txt)
        extr.config(text = self.extr_txt)
        chk.config(text = self.chk_txt)
        counter1.config(text = self.counter1_txt)
        chk2.config(text = self.chk2_txt)
        counter2.config(text = self.counter2_txt)
        com.config(text = self.com_txt)
        ent.config(text = self.ent_txt)
        chk3.config(text = self.chk3_txt)
        windowx.after(1000, self.update, err, ldat, extr, chk, counter1, chk2, counter2, com, ent, chk3, windowx)
        
    def load_data(self, filename, input_sheet, column_names, start_row, max_row):
        #print('Now translating' + str(filename))
        modulepath = os.path.dirname(xw.__file__)
        #print(modulepath)
        
        col = ','.join(column_names) #joins column names as a string with , - needed to set the columns for the dataframe
        rows = int(int(max_row) - int(start_row)) #number of rows is max_row minus start row - b/c max row refers to all of the filled 
        self.rows = rows
        
        self.ldat_txt = 'The length of data is ' + str(rows) + ' rows. The starting row is ' + str(start_row) +'\n'
        
        skipline = int(start_row) - 1 #this is done b/c pd starts with 0-index so 0 = 1st line, etc. and we want to skip the lines before the start row (so if starting row is 5 we want to skip 4 lines (0,1,2,3)
        data = pd.read_excel(filename, input_sheet, header=None, names=column_names, usecols=col, skiprows = skipline, nrows = rows) #looked at pandas documentation for how to work with excel
        
        if len(column_names) > 1:
            data[column_names] = data[column_names].astype(str) + '. '
            data["combined"] = data.sum(axis=1)
            untranslated = data["combined"].tolist()
        else:
            untranslated = data[column_names[0]].tolist()
        self.untranslated = untranslated
        
        names = pd.read_excel(filename, input_sheet, header=1, usecols = col) #creates a separate dataframe for the names of columns (which are not grabbed in the original dataframe due to the starting row being at the data, rather than the columns
        self.names = names #sets it so can be referenced outside of the function
        #print(names.columns) #as a check - it works
        self.extr_txt = 'Data extracted'
        
        #print(self.untranslated)
        
    def data_check(self, ilang, olang, filetype, column_names, input_column, filename, input_sheet, start_row, max_row):
        '''
        This function is to check the untranslated data, both for spelling using spellcheck function and against the dictionary created in create_glossary (thesaurus of CAAL terms) 
        so that terms which appear in the Russian glossary (keys) are replaced by the English values in each string. This is done to avoid inconsistent Google Translate translations, 
        which can sometimes translate one word into different things, particularly specialist terms such as the archaeological ones used here. As such, the glossary was made to allow for consistent tnanslation.
        
        The list of untranslated strings is iterated through, with empty lists for the obtained patterns and replacements. Each key in the glossary 
        is iterated through, with '*' wildcard characters added to the key, and the '[]' characters removed from the replacement string (as a str of the value would return it within brackets)
        Fnmatch is used to match the untranslated string un with the set pattern - I used fnmatch over re because it requires less specification, and it returns a boolean value which allows the check.
        If the value is True, i.e. if the pattern string matches up with the untranslated string, the key and replacement strings are added to a list for replacement. Once the glossary is iterated through,
        the reps and pats lists should contain an equal number of the words/phrases and desired replacements, which are then iterated through to replace within the string.
        
        The checked string (newstring) is then recapitalised according to which column it is in (if it is a Name column, all the individual words are capitalised - if anything else, the string is split up according to sentences and recapitalised)
        and then added to the checked_untranslated list.
        
        '''
        #print(str(ilang) + ' ' + str(olang))
        spell = Spelling()
        
        if str(ilang) == 'ru' and str(olang) == 'en':
            gloss = spell.create_glossary(ilang)
        elif str(ilang) == 'zh_CN' and str(olang) == 'en':
            gloss = spell.create_glossary(ilang)
        elif str(olang) == 'ru' and str(ilang) == 'en':
            gloss = spell.create_glossary(olang)
        elif str(olang) == 'zh_CN' and str(ilang) == 'en':
            gloss = spell.create_glossary(olang)
        else:
            gloss = {'1':['1']}
        
        self.chk_txt = 'Checking data for spelling errors and proper nouns'
        
        keep_capital = ['основное имя', 'название', 'aвтор', 'аддрес', 'Primary Name', 'Primary Address'] #defines list of columns to keep capitalised - includes 'Main Name', 'Name', 'Author', 'Address' columns which may be in various spreadsheets
        '''
        this is needed to determine if all the words in the string need to be capitalised after data check, which lowercases the whole string. If the string is part of the names column, every word needs to be capitalised
        so that the names of places are properly formatted. If the string is not part of names/addresses i.e. contains multiple sentences, then it will be split up into sentence-strings and each of those will be capitalised
        as I have yet to find a method to restore capital letters as they were.
        '''
        for col in self.names.columns: #for each column in the names dataframe created in input_data
            if any(c in col.lower() for c in keep_capital): #if a part of the lowercase column name matches any string in the keep_capital list
                capitalise = True #then capitalise is set to True
            else: #otherwise capitalise is set to False
                capitalise = False
        
        #print(capitalise) to check that it is registering as True/False - it is
        n=0    
        wb = xw.Book(filename) #defines wb as calling Book method for filename in xw
        sheet=wb.sheets[input_sheet] #sheet accesses the sheets from the workbook (using the input_sheet) #allows user to input column (varies based on
        #query and spreadsheet)   
        for un in self.untranslated: #for each string (un) in untranslated
            if str(un) == '' or str(un) == None:
                self.checked_untrans.append(un)
                n=n+1
                self.counter1_txt = str(n) + ' entries checked'
            elif str(un) == 'nan':
                self.checked_untrans.append('N/A')
            else:
                
                reps = []
                pats = [] #creates lists for replacements and patterns that are found through the check against the glossary
                
                check_un = spell.spell_checker(filetype, ilang, un, gloss)
                #print('working on it')
                if str(ilang) == 'ru' and str(olang) == 'en':
                    for key in gloss: #for each key in the russian glossary dictionary - so for each Russian term
                        pattern = '*' + str(key).lower() + '*' #the pattern is * wildcard character + the key + * - so it matches wherever the key may be present in the string. Fnmatch is used so * is accepted as a wildcard
                        replacement = str(gloss[key]).lower()[2:-2] #the replacement string which contains the English translation of the Russian key - lowercased as the whole string will be
                        if fn.fnmatch(check_un, str(pattern)) == True: #fnmatch automatically casematches, so lowers un and the pattern - if there is a match, so if boolean = True
                            reps.append(str(replacement)) #then append the replacement value to reps
                            pats.append(str(key).lower()) #and append the Russian word (which has a match) to pats - so the two still correspond with each other by index - lowercased as the rest of the string will be so if it isn't it won't get a match
                            continue
                        else: #otherwise reset the loop and continue
                            continue
                    
                    for i in range(0, int(int(len(pats)) - 1)): #if i = 0 to the length of pats (-1 because len starts from 1 while index starts from 0)
                        check_un = check_un.lower().replace(pats[i], reps[i]) #un becomes the lowercased string, with the pattern replaced
                    
                elif str(olang) == 'ru' and str(ilang) == 'en':
                    
                    for key in gloss:
                        pattern = '*' + str(gloss[key]).lower() + '*'
                        replacement = str(key).lower()[2:-2]
                        if fn.fnmatch(check_un.lower(), str(pattern)) == True: #fnmatch automatically casematches, so lowers un and the pattern - if there is a match, so if boolean = True
                            reps.append(str(replacement)) #then append the replacement value to reps
                            pats.append(str(gloss[key]).lower()) #and append the Russian word (which has a match) to pats - so the two still correspond with each other by index - lowercased as the rest of the string will be so if it isn't it won't get a match
                            continue
                        else: #otherwise reset the loop and continue
                            continue
                    for i in range(0, int(len(pats) - 1)): #if i = 0 to the length of pats (-1 because len starts from 1 while index starts from 0)
                        check_un = check_un.lower().replace(pats[i], reps[i]) #un becomes the lowercased string, with the pattern replaced
                elif str(olang) == 'zh_CN' and str(ilang) == 'en':
                    for key in gloss: #for each key in the russian glossary dictionary - so for each Russian term
                        pattern = '*' + str(key).lower() + '*' #the pattern is * wildcard character + the key + * - so it matches wherever the key may be present in the string. Fnmatch is used so * is accepted as a wildcard
                        replacement = str(gloss[key]).lower()[2:-2] #the replacement string which contains the English translation of the Russian key - lowercased as the whole string will be
                        if fn.fnmatch(check_un, str(pattern)) == True: #fnmatch automatically casematches, so lowers un and the pattern - if there is a match, so if boolean = True
                            reps.append(str(replacement)) #then append the replacement value to reps
                            pats.append(str(key).lower()) #and append the Russian word (which has a match) to pats - so the two still correspond with each other by index - lowercased as the rest of the string will be so if it isn't it won't get a match
                            continue
                        else: #otherwise reset the loop and continue
                            continue
                    
                    for i in range(0, int(int(len(pats)) - 1)): #if i = 0 to the length of pats (-1 because len starts from 1 while index starts from 0)
                        check_un = check_un.lower().replace(pats[i], reps[i]) #un becomes the lowercased string, with the pattern replaced
                elif str(ilang) == 'zh_CN' and str(olang) == 'en':
                    for key in gloss:
                        pattern = '*' + str(gloss[key]).lower() + '*'
                        replacement = str(key).lower()[2:-2]
                        if fn.fnmatch(check_un.lower(), str(pattern)) == True: #fnmatch automatically casematches, so lowers un and the pattern - if there is a match, so if boolean = True
                            reps.append(str(replacement)) #then append the replacement value to reps
                            pats.append(str(gloss[key]).lower()) #and append the Russian word (which has a match) to pats - so the two still correspond with each other by index - lowercased as the rest of the string will be so if it isn't it won't get a match
                            continue
                        else: #otherwise reset the loop and continue
                            continue
                else:
                    pass
                
                
                nu = check_un.replace('.,', ',') #removes duplicated punctuation which can appear after the data check
                newstring = nu.replace('..', '.') #removes duplicated punctuation appearing after data check
    
                '''
                   The following method is somewhat finicky and could be improved on, but involves the recapitalisation of the string after it was made all-lowercase to match properly with the glossary terms
                   It is finicky because there are multiple sentences within strings, so .capitalize() on the whole string doesn't work because it uppercases the first letter but makes the rest of the characters lowercase - we want each
                   sentence, as well as names, capitalised. So instead the string is split up along '.' which encapsulates both sentences and names when initials and last names are given (as is the standard in these
                   descriptions) and then capitalized, before being joined back together and appended. There is also a very specific method intended to capitalise the initials as well by checking if the last word of 
                   the list of words in a sentence is only 1 character (which should only be with initials) and capitalising that
                   I will continue looking for ways to improve on this but this is the most thorough one I have come up with.
                '''
                if capitalise == True: #if capitalize is set to True from earlier
                    caps = newstring.strip().title()#then caps is the newstring with every word capitalised (.title())
                    #print(newstring)
                    self.checked_untrans.append(caps) #and appended to checked_untrans list for translation
                    n=n+1
                    self.counter1_txt = str(n) + ' entries checked'
                else: #otherwise (if capitalise is False)
                    nu = newstring.split('.') #splits newstring along '.' as separators so it's individual sentences
                    cap = [] #creates empty list to store these in
                   
                    for sent in nu: #for each sentence in the nu string
                        caps = sent.strip().capitalize() #strips leading/tailing spaces
                        sep = caps.split(' ') #splits this string along ' ' - so now a consists of all the individual words within sentence z - this is done so initials are capitalised as they end up being lowercase 
                        #because the above-splitter is along '.' so things like 'V.Katsev' will become split
                        last = len(str(sep[-1])) #v is the length of the last word within the list of words in a
                        #if the last word is only one character (which if Russian grammar/spelling rules are followed should only be in the case of initials i.e. 'house of v' and 'Katsev' are separate when it should be house of V.Katsev
                        if last == 1: 
                            capp = sep[-1].capitalize() #then the last word is capitalised
                            sep.pop(-1) #the previous last word is removed
                            sep.append(capp) #and the new capitalised initial is added instead
                            joined = ' '.join(sep) #the string is joined back together using ' ' as separator
                            cap.append(joined) #and appended
                        elif last > 1: #otherwise if the last word is more than one character
                            cap.append(caps) #append the capitalised sentence as it was
                           
                    fin = '. '.join(cap) #rejoins the sentences in cap with a period 
                    
                    #print('1.' + un + '\n 2.' + check_un + '\n 3.' + fin) #check
                    self.checked_untrans.append(fin) #appends to checked untranslated
                    n=n+1
                    self.counter1_txt = str(n) + ' entries checked'
                    
                    translated = self.translator(ilang, olang, fin, n)
                    translated2 = self.combinedata(column_names, input_column, translated, fin)
                    self.input_data(filename, input_sheet, input_column, start_row, translated2, wb, sheet, n)
        wb.save()
        wb.close()
                    
                
    def translator (self, ilanguage, olanguage, i, n):
        
        self.chk2_txt = 'Now translating - this can take a while, make a cup of tea :)'
        try:
            #n=0
            #for i in self.checked_untrans: #for strings in untranslated
                
            transl_text = str(i) #converts it into a string (to be sure that it is actually coming out as string value)
            transltxt = transl(source=ilanguage, target=olanguage).translate(transl_text) #translates the string - taken from deep_translator documentation#
            #self.translated.append(transltxt)
            self.counter2_txt = str(n) + ' entries translated'
            return transltxt
            #so progress can be checked - and can see if it gets stuck on any particular entries
            #print(transltxt) #optional - can uncheck to see the translated results
            #print('translating2')
            #print(i + '\n \n' + transltxt +'\n \n')
        except Exception as e: #except if there is an error 
            self.err_txt = str(traceback.print_exc())
            return 'could not translate'
            #print(str(traceback.print_exc()))
            
            #print('''IF THE CONNECTION HAS TIMED OUT, WAIT SOME TIME/TRY AGAIN OR TRY PUTTING IN headers={
                                #'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_9_3) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/35.0.1916.47 Safari/537.36'}
                                # INTO deep translator's google_trans translate function in response = requests.get on line 117 ''')
 
        #print("If it's not proceeding, check if Excel has opened the file because it contains macros, click 'x' on the macros to access the spreadsheet and program will resume")    
            
    def combinedata(self, column_names, input_column, translated, untranslated):
        '''
        This function is for columns like Location Notes, where the English translation is input in the same entry underneath the Russian data rather than in its own column. So we need to combine
        the untranslated and translated data into one entry, which can then be input into the column
        '''
        combine = False
        self.combine = combine
        if input_column in column_names: #if the input_column is also in column_names aka if data is being extracted from the same column that it should be put back into - so there is both Russian and English in the entry
            self.com_txt = 'Combining data from ' + str(column_names)
            
            combine = True #so if we want to be inputting the result to the same column that the original data was in without overriding the original data
            self.combine = combine #sets it so it can be referenced outside of the method
            
            if str(translated).lower() == 'nan':
                #self.combineddata.append('N/A')
                return 'N/A'
            else:
                #num = int(self.untranslated.index(i))
                combined = str(untranslated) + ' \n ' + str(translated)
                self.combineddata.append(combined) #appends combined string to combinedddata
                return combined
                #print(combined)
                
            
        else:
            combine = False #combine is set to false
            self.combine = combine #sets it so we can reference it outside of this method
            #pass
            return translated

        
    def input_data(self, filename, input_sheet, input_column, start_row, translated, wb, sheet, n):

        '''
        This function inputs the translated data from the list into the cells using xlwings as xw - it opens the Sheet, inputs the data using a combination of the input_column and given row, knowing that the data is in
        the same order because that's the order it was extracted in and worked on.
        
        The below code emerged from combined xlwings documentation and various tidbits of code obtained from StackOverflow as I was running into errors.
        '''
        
        #wb = xw.Book(filename) #defines wb as calling Book method for filename in xw
        #sheet=wb.sheets[input_sheet] #sheet accesses the sheets from the workbook (using the input_sheet) #allows user to input column (varies based on
        #query and spreadsheet)
        self.ent_txt = 'Entering entry ' + str(n) + ' into spreadsheet'
        
        #for row in range(int(start_row), int(len(self.translated)) + 1): #for row in the range of starting row to max_row
        row = int(start_row) + int(n) - 1
        cell = str(input_column + str(row)) #the cell is the string combination of input_column and row - need to convert row to string so it can be concatenated with column string
            #allowing you to define which cell the data is going in
        print(cell)
        sheet.range(cell).value = str(translated)
        '''
        if self.combine == False:
            sheet.range(cell).value = self.translated[i] #the xw command for writing into a cell where value is defined as i in translated
        else:
            sheet.range(cell).value = self.combineddata[i]
        
            i = i + 1 #i adds 1 to iterate through the list ''' 
        self.chk3_txt = str(n) + ' entries input'
                #this prevents the already-inputted cell values from being overwritten by future iterations
                
        wb.save() #saves the changes
        #wb.close() #closes the workbook to allow next columns to be translated
        
    '''def sql_work(self, filename, folder, input_sheet, tables, data_cols, max_row):
        ''''''
        This module takes the translated data from the file we are working in (called after input_data), and extracts it to a dataframe. Each row in the dataframe is tehn added
        ''''''
        slite = SQLite(folder)
        
        data = pd.read_excel(filename, input_sheet, header=0, usecols=data_cols, nrows = self.rows)
        data.to_csv('SQLdata.csv')
        csvname = 'SQLdata.csv'
        pn = os.path.split(os.path.dirname(os.path.abspath(__file__)))[0] 
        
        csvfile = open(os.path.join(pn, 'src', csvname), 'r',encoding="UTF-8")''''''csvfile defined as open the joined path of data_folder and the csvname''''''
        reader = csv.reader(csvfile)
    
        for t in tables:
            if str(t[ :7]).lower() in str(filename).lower():
                for i in reader: #for each row (i) in the csvfile "Seals.csv" (reader)
                    #skip the heading row
                    if reader.line_num == 1:
                        continue
                    slite.insert_row(t, i)    
            else:
                continue'''
        
            
    def reset(self): #resets the function so there are no leftover values in the lists for the future values
        self.untranslated.clear() #clears the untranslated list
        self.checked_untrans.clear() #ibid for checked untranslated
        self.translated.clear() #ibid for translated
        self.combineddata.clear()
        self.ldat_txt = '...'
        self.extr_txt = '...'
        self.chk_txt = '...'
        self.counter1_txt = '...'
        self.chk2_txt = '...'
        self.counter2_txt = '...'
        self.com_txt = '...'
        self.ent_txt = '...'
        self.chk3_txt = '...'
       
    
class TranslateRun: #the class to run the above methods
    transl = TranslMethods() #transl is the TranslMethods class
        
    def runn(self, filename, input_sheet, column_names, input_column, start_row, max_row, ilanguage, olanguage, filetype): #the whole translation and SQL input process with the associated variables
        self.transl.load_data(filename, input_sheet, column_names, start_row, max_row)
        self.transl.data_check(ilanguage, olanguage, filetype, column_names, input_column, filename, input_sheet, start_row, max_row)
        #self.transl.translator(ilanguage, olanguage)
        #self.transl.combinedata(column_names, input_column)
        #self.transl.input_data(filename, input_sheet, input_column, start_row, max_row)
    
    '''def SQLinput(self, filename, folder, input_sheet, tables, data_cols, max_row):
        self.transl.sql_work(filename, folder, input_sheet, tables, data_cols, max_row)'''
        
    
    def runtransl(self, filename, input_sheet, column_names, input_column, start_row, max_row, ilanguage, olanguage, windowx, filetype):

        framea = tk.Frame(master=windowx, width = 150, height=100, padx=5, pady=5)
        framea.pack(fill = tk.X, side=tk.TOP)
        
        frameb = tk.Frame(master=windowx, width = 150, height = 100, padx=5, pady=5)
        frameb.pack(fill=tk.X, side = tk.TOP)
        
        framec = tk.Frame(master=windowx, width = 150, height = 100, padx=5, pady=5)
        framec.pack(fill=tk.X, side = tk.TOP) 
        
        fl = tk.Label(master=framea, text = 'Now translating: ' + str(filename) + '\n Columns: ' + str(column_names) +'-->' + str(input_column))
        fl.pack(fill = tk.X, side = tk.TOP)
        
        note = tk.Label(master=framea, text = 'Note: If file contains forms/macros you will need to manually close/log into these when they pop-up during entry')
        note.pack(fill = tk.X, side = tk.TOP)
        
        err = tk.Label(master=framea, text = '...')
        err.pack(fill=tk.X, side = tk.TOP)
        
        ldat = tk.Label(master=framea, text = '...')
        ldat.pack(fill = tk.X, side = tk.TOP)
    
        extr = tk.Label(master=framea, text = '...')
        extr.pack(fill = tk.X, side = tk.BOTTOM)
        
        chk = tk.Label(master=frameb, text = '...')
        chk.pack(fill=tk.X, side = tk.TOP)
        
        counter1 = tk.Label(master=frameb, text = '...')
        counter1.pack(fill=tk.X, side = tk.BOTTOM)
        
        framed = tk.Frame(master=windowx, width = 150, height=100, padx=5, pady=5)
        framed.pack(fill = tk.X, side=tk.TOP)
        
        framee = tk.Frame(master=windowx, width = 150, height=100, padx=5, pady=5)
        framee.pack(fill = tk.X, side=tk.TOP)
        
        framef = tk.Frame(master=windowx, width = 150, height=100, padx=5, pady=5)
        framef.pack(fill = tk.X, side=tk.TOP)
        
        chk2 = tk.Label(master=framec, text='...')
        chk2.pack(fill=tk.X, side = tk.TOP)
             
        counter2 = tk.Label(master=framec, text = '...')
        counter2.pack(fill=tk.X, side = tk.BOTTOM)
        
        com = tk.Label(master=framed, text = '...')
        com.pack(fill=tk.X, side = tk.BOTTOM)
    
        ent = tk.Label(master=framee, text = '...')
        ent.pack(fill = tk.X, side = tk.TOP)
        
        chk3 = tk.Label(master=framee, text = '...')
        chk3.pack(fill = tk.X, side = tk.BOTTOM)
        
        
        windowx.after(1000, self.transl.update, err, ldat, extr, chk, counter1, chk2, counter2, com, ent, chk3, windowx)
        self.runn(filename, input_sheet, column_names, input_column, start_row, max_row, ilanguage, olanguage, filetype)
        framea.forget()
        frameb.forget()
        framec.forget()
        framed.forget()
        framee.forget()
        framef.forget()
        self.transl.reset()
    
class TranslateDoc:
    pdf_content = []
    doc = []
    tables = []
    translated_table = []
    row_length = []
    column_length = []
    translated_doc = []
    err_txt = ' '
    counter0_txt = '...'
    ldat_txt = '...'
    counter1_txt = '...'
    counter2_txt = '...'
    ent_txt = '...'
    chk3_txt = '...'
    resett = False
    
    def update(self, err, counter0, ldat, counter1, counter2, ent, chk3, window): #function to allow window to continue updating
        err.config(text = self.err_txt)
        counter0.config(text = self.counter0_txt)
        ldat.config(text = self.ldat_txt)
        counter1.config(text = self.counter1_txt)
        counter2.config(text = self.counter2_txt)
        ent.config(text = self.ent_txt)
        chk3.config(text = self.chk3_txt)
        window.after(2000, self.update, err, counter0, ldat, counter1, counter2, ent, chk3, window)
        
    def load_data(self, filename, filetype, ilang):
        
        if filetype == 'Word':
                untrans = docx.Document(filename)
                self.untrans = untrans
                self.ldat_txt = 'Text extracted from file'
        elif filetype == 'PDF':
            self.ldat_txt = "This may take quite a long time, don't worry"
            pdf_file = open(filename, 'rb')
            read_pdf = PyPDF2.PdfFileReader(filename)
            
            #get the pages
            number_of_pages = read_pdf.getNumPages()
            
            
            #iterate from the pages of the pdf
            for i in range(number_of_pages):
                
                #get the page
                page = read_pdf.getPage(i)
                
                #get the text and put everything to lower case
                page_content = page.extractText()
                
                #print content to see it to make sure it is reading the file(note we can remove this)
                #print(page_content)

        
            #add the total document text
                self.pdf_content.append(page_content)
                
                n=0
                if page_content == None or page_content == '':
                    try:
                        pn = os.path.split(os.path.dirname(os.path.abspath(__file__)))[0]
                        poppler = os.path.join(pn, 'data', 'poppler-0.68.0', 'bin')
            
                        #get it to read PDF extract text as object translate it and output as a doc
                        #read_pdf = PyPDF2.PdfFileReader(filename)
                        #page_number = threading.Thread(target=read_pdf.getNumPages()).start
                        #print(poppler)
                        read_pdf = PyPDF2.PdfFileReader(filename)
                        
                        PIL.Image.MAX_IMAGE_PIXELS = 1500000000000
                        
                        images = convert_from_path(filename, poppler_path =poppler, grayscale=True)
                        self.counter0_txt = 'Extracted pages from PDF'
                        tess = os.path.join(pn, 'TesseractOCR', 'tesseract.exe')
                        pytesseract.pytesseract.tesseract_cmd = tess
                    
                    
                        if ilang == 'ru':
                            lan = 'rus'
                        elif ilang == 'en':
                            lan = 'eng'
                        elif ilang == 'kk':
                            lan = 'kaz'
                        elif ilang == 'ky':
                            lan = 'kir'
                        elif ilang == 'fr':
                            lan = 'fra'
                        elif ilang == 'ar':
                            lan = 'ara'
                        elif ilang == 'de':
                            lan = 'deu'
                        elif ilang == 'zh_CN':
                            lan = 'chi_sim'
                        elif ilang == 'es':
                            lan = 'spa'
                        elif ilang == 'uz':
                            lan = 'uzb'
                        elif ilang == 'tg':
                            lan = 'tgk'
                        elif ilang == 'it':
                            lan = 'ita'
                        else:
                            lan = 'eng'
                        
                        for img in enumerate(images):
                            txt = pytesseract.image_to_string(img, lang = lan)
                            self.pdf_content.append(str(txt))
                            #print(str(txt))
                            n+=1
                            self.counter0_txt = str(n) + ' pages extracted from PDF'
                    except Exception as e:
                        self.err.update = str(e.args())
                        self.untrans.append('Could not process this page')
                        continue
                else:
                    n+=1
                    self.counter0_txt = str(n) + ' pages of text scraped from PDF'
            
            #close the pdf
            pdf_file.close()
            
            
            
            
            '''
            for page, page_data in enumerate(convert):
                txt = pytesseract.image_to_string(page_data, lang = lan)
                pg_txt =("Page # {} - {}".format(str(page), txt))
                self.pdf_content.append(pg_txt)
                n+=1
                print(str(pg_txt))
                self.counter0_txt = str(n) + ' pages processed from PDF'
            print(str(self.pdf_content) + 'done')
        '''
                        
            self.ldat_txt = 'Text extracted from pages'
        else:
            pass
        
    def check_data(self, filetype, ilang, olang, window):
        
        spell = Spelling()
        if ilang == 'ru' and olang == 'en':
            gloss = spell.create_glossary(ilang)
        elif ilang == 'zh_CN' and olang == 'en':
            gloss = spell.create_glossary(ilang)
        if olang == 'ru' and ilang == 'en':
            gloss = spell.create_glossary(olang)
        if olang == 'zh_CN' and ilang == 'en':
            gloss = spell.create_glossary(olang)
        else:
            gloss = {'1':['1']}
            
        
        if filetype == 'Word':
                n=0
                if len(self.untrans.tables) > 0:
                    for table in self.untrans.tables:
                        transl_table = []
                        self.row_length.append(len(table.rows))
                        #print(self.row_length)
                        self.column_length.append(len(table.columns))
                        #print(self.column_length)
                        for row in table.rows:
                            for cell in row.cells:
                                cell_text = []
                                for para in cell.paragraphs:
                                    parag = str(para.text)
                                    checked_table_parag = self.check_data2(filetype, ilang, olang, parag, gloss)
                                    cell_text.append(checked_table_parag)
                                nu_cell = '\n'.join(cell_text)
                                #itable.append(nu_cell)
                                #print(nu_cell)
                                self.tables.append(nu_cell)
                                transl_cell = self.translate_data(ilang, olang, nu_cell)
                                #print(transl_cell)
                                transl_table.append(transl_cell)
                            n = n+1
                            self.counter1_txt = str(n) + ' rows checked out of ' + str(len(table.rows)) + ' from ' + str(len(self.untrans.tables)) + ' table(s)'
                            window.update()
                        self.translated_table.append(transl_table)
                        
                                    
                else:
                    pass
                
                if len(self.untrans.paragraphs) > 0:
                    for para in self.untrans.paragraphs:
                        parag = str(para.text)
                        checked_parag = self.check_data2(filetype, ilang, olang, parag, gloss)
                        self.doc.append(checked_parag)
                        n = n + 1
                        self.counter1_txt = str(n) + ' paragraphs checked out of ' + str(len(self.untrans.paragraphs))
                    
                    for parag in self.doc:
                        transl_parag = self.translate_data(ilang, olang, parag)
                        self.translated_doc.append(transl_parag)
                        #print(transl_parag)
                        
        elif filetype == 'PDF':
            for page in self.pdf_content:
                nu_page = str(page).replace('\n', ' ')
                #checked_page = self.check_data2(filetype, ilang, olang, str(nu_page), gloss)
                transl_page = self.translate_data(ilang, olang, nu_page)
                #print(transl_page)
                self.translated_doc.append(transl_page)
        else:
            pass
        
    def check_data2(self, filetype, ilang, olang, parag, gloss):
        spell = Spelling()
        for word in parag.split(' '): #splits paragraph into individual words
            
            punc_free = re.sub(r'(?s:[,.!:;"?])\Z', '', word) #removes punctuation from word 
            #print(punc_free)
            #print(punc_free)
            if len(str(punc_free)) == 0 or str(punc_free) == None:
                continue
            elif len(str(punc_free)) > 1 and str(punc_free).strip()[0].isupper(): #if it's a capitalised word skip checking because likely proper noun
                continue
            else:
                checked_word = spell.spell_checker(filetype, ilang, punc_free, gloss)
                if checked_word == None:
                    continue
                else:
                    if ilang == 'ru' and olang == 'en':
                        for key in gloss:
                            try:
                                pattern = str(key) + '*' #the pattern is * wildcard character + the key + * - so it matches wherever the key may be present in the string. Fnmatch is used so * is accepted as a wildcard
                                replacement = str(gloss[key]).lower()[2:-2] #the replacement string which contains the English translation of the Russian key - lowercased as the whole string will be
                                if fn.fnmatch(word, str(pattern)) == True: #fnmatch automatically casematches, so lowers un and the pattern - if there is a match, so if boolean = True
                                    parag = parag.replace(punc_free, replacement)
                                elif fn.fnmatch(checked_word, str(pattern)) == True:
                                    parag = parag.replace(punc_free, replacement)
                                else: #otherwise reset the loop and continue
                                    parag = parag.replace(punc_free, checked_word)
                                    
                            except Exception as e:
                                self.err_txt = str(e.args) + ', skipping word'
                                parag = parag.replace(punc_free, checked_word)
                                continue
                    elif ilang == 'en' and olang =='zh_CN':
                        for key in gloss:
                            try:
                                pattern = str(key) + '*' #the pattern is * wildcard character + the key + * - so it matches wherever the key may be present in the string. Fnmatch is used so * is accepted as a wildcard
                                replacement = str(gloss[key]).lower()[2:-2] #the replacement string which contains the English translation of the Russian key - lowercased as the whole string will be
                                if fn.fnmatch(word, str(pattern)) == True: #fnmatch automatically casematches, so lowers un and the pattern - if there is a match, so if boolean = True
                                    parag = parag.replace(punc_free, replacement)
                                elif fn.fnmatch(checked_word, str(pattern)) == True:
                                    parag = parag.replace(punc_free, replacement)
                                else: #otherwise reset the loop and continue
                                    parag = parag.replace(punc_free, checked_word)
                                    
                            except Exception as e:
                                self.err_txt = str(e.args) + ', skipping word'
                                parag = parag.replace(punc_free, checked_word)
                                continue
                    elif olang == 'ru' and ilang == 'en':
                        for key in gloss:
                            try:
                                pattern = '*' + str(gloss[key]) + '*'
                                replacement = str(key).lower()[2:-2]
                                if fn.fnmatch(word, str(pattern)) == True: #fnmatch automatically casematches, so lowers un and the pattern - if there is a match, so if boolean = True
                                    parag = parag.replace(punc_free, replacement)
                                elif fn.fnmatch(checked_word, str(pattern)) == True:
                                    parag = parag.replace(punc_free, replacement)
                                else: #otherwise reset the loop and continue
                                    parag = parag.replace(punc_free, checked_word)
                                        
                            except Exception as e:
                                self.err_txt = str(e.args) + ', skipping word'
                                parag = parag.replace(punc_free, checked_word)
                                continue
                    elif olang == 'en' and ilang == 'zh_CN':
                        for key in gloss:
                            try:
                                pattern = '*' + str(gloss[key]) + '*'
                                replacement = str(key).lower()[2:-2]
                                if fn.fnmatch(word, str(pattern)) == True: #fnmatch automatically casematches, so lowers un and the pattern - if there is a match, so if boolean = True
                                    parag = parag.replace(punc_free, replacement)
                                elif fn.fnmatch(checked_word, str(pattern)) == True:
                                    parag = parag.replace(punc_free, replacement)
                                else: #otherwise reset the loop and continue
                                    parag = parag.replace(punc_free, checked_word)
                                        
                            except Exception as e:
                                self.err_txt = str(e.args) + ', skipping word'
                                parag = parag.replace(punc_free, checked_word)
                                continue
                    else:
                        parag = parag.replace(punc_free, checked_word)
                    
        
        return parag           
            
    def translate_data(self, ilanguage, olanguage, paragraph):
        try:
                transl_text = str(paragraph) #converts it into a string (to be sure that it is actually coming out as string value)
                transltxt = transl(source=ilanguage, target=olanguage).translate(transl_text) #translates the string - taken from deep_translator documentation
                #self.translated_doc.append(transltxt) #appends translated string to list in same order
                return transltxt
                #print(transltxt)
                
                #so progress can be checked - and can see if it gets stuck on any particular entries
                #print(transltxt) #optional - can uncheck to see the translated results
                #print('translating2')
     
        except Exception as e: #except if there is an error 
            self.err_txt = str(e.args)
            
            
    def write_file(self, newdoc, location, filename, olang, filetype):
        #print('location' + str(len(location)))
        langs_dict = transl.get_supported_languages(as_dict = True)
        lang_list = list(langs_dict.keys())
        abb_list = list(langs_dict.values())
        
        position = abb_list.index(olang)
        long_lang = lang_list[position]
        
        if newdoc == 'New document':
            ndoc = docx.Document()
            if len(str(location)) > 2:
                filen = os.path.basename(filename)
                nfile = filen[:-5] + '_translated_' + long_lang + '.docx'
                new_filename = os.path.join(location, nfile)
            else:
                if filetype == 'PDF':
                    
                    new_filename = str(filename)[:-3] + '_translated_' + long_lang + '.docx'
                else:
                    new_filename = str(filename)[:-5] + '_translated_' + long_lang + '.docx'
             
            
            file = r''+ new_filename
            self.ent_txt = 'New file created: ' + str(file)
            #print(file)
        
            if len(self.translated_table) > 0:
                n=0
                for itable in self.translated_table:
                    rowlength = self.row_length[self.translated_table.index(itable)]
                    columnlength = self.column_length[self.translated_table.index(itable)]
                    df = pd.DataFrame(np.array(itable).reshape(rowlength, columnlength))
                    word_table = ndoc.add_table(rows = rowlength,cols = columnlength)
                    for x in range(0, rowlength):
                        for y in range(0, columnlength):
                            cell = word_table.cell(x,y)
                            try:
                                cell.text = df.iloc[x,y]
                            except:
                                cell.text = 'Error inputting'
                                continue
                            ndoc.save(file)
            else:
                pass
            
            if len(self.translated_doc) > 0:
                n=0
                for parag in self.translated_doc:
                    ndoc.add_paragraph(parag)
                    ndoc.save(file)
                    n = n + 1
                    self.chk3_txt = str(n) + ' paragraphs entered out of ' + str(len(self.translated_doc))
            else:
                pass
            
        else:
            n=0
            ndoc = docx.Document(filename)
            ndoc.add_paragraph('\n \n' + 'Translated to ' + str(long_lang.capitalize()) + '\n')
            ndoc.save(filename)
            self.ent_txt = 'Entering into existing file'
            
            
            if len(self.translated_table) > 0:
                n=0
                for itable in self.translated_table:
                    rowlength = self.row_length[self.translated_table.index(itable)]
                    columnlength = self.column_length[self.translated_table.index(itable)] 
                    df = pd.DataFrame(np.array(itable).reshape(rowlength, columnlength))
                    word_table = ndoc.add_table(rows = rowlength,cols = columnlength)
                    for x in range(0, rowlength):
                        for y in range(0, columnlength):
                            cell = word_table.cell(x,y)
                            try:
                                cell.text = df.iloc[x,y]
                            except:
                                cell.text = 'Error inputting'
                                continue
                            ndoc.save(filename)
            else:
                pass
            
            if len(self.translated_doc) > 0:
                n=0
                for parag in self.translated_doc:
                    ndoc.add_paragraph(parag)
                    ndoc.save(filename)
                    n = n + 1
                    self.chk3_txt = str(n) + ' paragraphs entered out of ' + str(len(self.untrans.paragraphs))
            else:
                pass
            
    def reset(self): #resets the function so there are no leftover values in the lists for the future values
        self.pdf_content.clear()
        self.tables.clear()
        self.translated_table.clear()
        self.row_length.clear()
        self.column_length.clear()
        self.doc.clear() #clears the untranslated list
        self.translated_doc.clear() #ibid for checked untranslated
        self.ldat_txt = '...'
        self.counter1_txt = '...'
        self.counter2_txt = '...'
        self.ent_txt = '...'
        self.chk3_txt = '...'
        
    
class TranslDocRun:
    
    def run(self, window, filename, filetype, location, ilanguage, olanguage, newdoc):
        TD = TranslateDoc()
        framea=tk.Frame(master=window, width = 150, height = 100, padx=5, pady=5)
        framea.pack(fill=tk.X, side = tk.TOP)
        frameb = tk.Frame(master=window, width = 150, height = 100, padx=5, pady=5)
        fl = tk.Label(master=framea, text = 'Now translating: ' + str(filename))
        fl.pack(fill=tk.X, side = tk.TOP)
        framec = tk.Frame(master=window, width = 150, height = 100, padx=5, pady=5)
        framed = tk.Frame(master=window, width = 150, height=100, padx=5, pady=5)
        framee = tk.Frame(master=window, width = 150, height=100, padx=5, pady=5)
        
        
        err = tk.Label(master=framea, text = '...')
        counter0 = tk.Label(master=framea, text = '...')
        ldat = tk.Label(master=framea, text = '...')
        counter1 = tk.Label(master=frameb, text = '...')
        counter2 = tk.Label(master=framec, text = '...')
        ent = tk.Label(master=framed, text = '...')
        chk3 = tk.Label(master=framee, text = '...')
        
        frameb.pack(fill=tk.X, side = tk.TOP)
        self.frameb = frameb
        
        framec.pack(fill=tk.X, side = tk.TOP) 
        self.framec=framec
        
        err.pack(fill=tk.X, side = tk.TOP)
        self.err=err
        
        counter0.pack(fill = tk.X, side = tk.BOTTOM)
        self.counter0=counter0
        
        ldat.pack(fill = tk.X, side = tk.BOTTOM)
        self.ldat=ldat
        
        counter1.pack(fill=tk.X, side = tk.BOTTOM)
        self.counter1=counter1
        
        framed.pack(fill = tk.X, side=tk.TOP)
        self.framed=framed
        
        framee.pack(fill = tk.X, side=tk.TOP)
        self.framee=framee
        
        counter2.pack(fill=tk.X, side = tk.BOTTOM)
        self.counter2=counter2
        
        ent.pack(fill = tk.X, side = tk.TOP)
        self.ent=ent
        
        chk3.pack(fill = tk.X, side = tk.BOTTOM)
        self.chk3=chk3
        
        window.after(1000, TD.update, err, counter0, ldat, counter1, counter2, ent, chk3, window)
        
        try:
            TD.load_data(filename, filetype, ilanguage)
            window.update()
            TD.check_data(filetype, ilanguage, olanguage, window)
            window.update()
            TD.write_file(newdoc, location, filename, olanguage, filetype)
            window.update()
            framea.forget()
            frameb.forget()
            framec.forget()
            framed.forget()
            framee.forget()
            TD.reset()
        except Exception as e: 
            windoww = tk.Tk()
            fram = tk.Frame(master=windoww, width = 150, height = 150, padx = 5, pady=5)
            fram.pack(fill = tk.X, side = tk.TOP)
            err = tk.Label(master=fram, text = 'ERROR:' + str(e) + ' ' + str(e.args) + '/n' + str(traceback.print_exc(limit=None, file=None, chain=True)))
            err.pack(fill=tk.X, side = tk.TOP)
            tip = tk.Label(master=fram, text = 'If Permission Error check permissions on files and tick "allowed to edit" for everyone or check that files are not open \n If not translating or "connection error" wait some time and try again as this is due to API')
            tip.pack(fill=tk.X, side = tk.TOP)
            
            def resett():
                reset = ResetGUI()
                reset.restart()
                
            framee = tk.Frame(master=windoww, width = 150, height = 150, padx = 5, pady=5)
            framee.pack()
            new_btn = tk.Button(master=framee, text = 'RESTART', command = resett)
            new_btn.pack(fill = tk.X, side = tk.BOTTOM)
        
            windoww.mainloop()
        
        
class Spelling():
    newstring = None
    
    def ru_dict(self):
        '''
        A short function to move dictionary files from the data folder in src code to the pyenchant folder
        on the user's drive so that pyenchant can then use it as a Russian dictionary. This is done because although
        I could move the files manually to my own pyenchant folder, it needs to also be done for any user. 
        It is simply done by finding the paths to the dictionary files and to the module, and then renaming the file path 
        which moves the files to the required directory. The files come from LibreOffice and are open access, I was able 
        to get them by downloading the software and looking in its language folders using advice from a Russian article
        on how to write a spell-checker.
        '''
        pn = os.path.split(os.path.dirname(os.path.abspath(__file__)))[0] 
        data = os.path.join(pn, 'data', 'dict', 'dictionary files')
        files = os.listdir(data)
        
        modulepath = os.path.dirname(enchant.__file__)
        folder = os.path.join(modulepath, 'data', 'mingw64', 'share', 'enchant','hunspell')
        #print(folder)
        for file in files:
            file_path = os.path.join(data, file)
            new_path = os.path.join(folder,file)
            if os.path.exists(new_path):
                pass
            else:
                os.rename(file_path, new_path)
    
    def create_glossary(self, ilang): 
        
        '''
        here we need to create a dictionary which the untranslated data can be compared against using the thesaurus spreadsheet developed by CAAL. The ditionary is created by using pandas to read the spreadsheet
        and make a dataframe of it, converting it to a CSV file as for some reason the Russian here did not want to encode otheriwse, and creating a dictionary with the Russian words set as keys and the English words
        as values; this means that strings can then be compared against the keys to see if they are present, and then substitute the keys for the values.
        
        '''
        
        project_folder = os.path.split(os.path.dirname(os.path.abspath(__file__)))[0] #splits path from the directory name to the main module being opened here
        data_folder = os.path.join(project_folder) #joins the main module path with that of data folder so it can be accessed
        glossfile = os.path.join(data_folder, 'data', "glossary.xlsx") #joins the data folder path with image folder so it can be opened/accessed
        
        #various StackOverflow queries helped with understanding errors that were arising here, particularly in terms of encoding
        if ilang == 'ru':
            
            rus = pd.read_excel(glossfile, usecols = 'B:C')#uses pandas to read excel file extracting columns B to C (Russian and English)
            rus.to_csv("Rusgloss.csv", index = None, header=True) #converts excel to csv because otherwise there are encoding issues that cannot be addressed in pandas read_excel because it removed ability to specify encoding
            rusdf = pd.DataFrame(pd.read_csv("Rusgloss.csv", encoding = 'utf-8')) #reads created csvfile to dataframe specifying the encoding so the Russian characters are read properly
            russgloss = rusdf.set_index('Russian').T.to_dict('list')
            ##print(russgloss)
            self.russgloss = russgloss #sets it so it can be referred
            return russgloss
            #print('Glossary created')
        elif ilang == 'zh_CN':
            cn = pd.read_excel(glossfile, usecols = 'C:D')
            cn.to_csv('Chinese_gloss.csv', index = None, header = True)
            cndf = pd.dataframe(pd.read_csv('Russgloss.csv', encoding = 'utf-8'))
            cngloss = cndf.set_index('Chinese').T.to_dict('list')
            self.cngloss = cngloss
            return cngloss
        
    def xl_spellcheck(self, string, ilang):
        '''A function to check each word for misspellings by comparing it with the pyenchant spellchecker using the russian dictionary files as d. 
        This function is applied to each untranslated string within translmethods.
        To operate the spell-checker, we first need to remove proper nouns as they will inevitably come back as misspelled and be erroneously replaced. This is done by using re to strip the string
        of any words beginning with a capital letter - which I know will also omit nouns at the start of a sentence so a better method will be implemented in the future.
        Then words in Russian quotes are also removed as these relate to names or Central Asian terms which are better off not being corrected in the first place
        And then punctuation . and , are removed to make analysis easier
        And double spaces are set to single spaces and trailing spaces are removed again '''
        
        global newstring #makes newstring a global variable so it can be accessed from translate methods
        newstring = string
        
        d = enchant.Dict('en_GB')
        
        pn = os.path.split(os.path.dirname(os.path.abspath(__file__)))[0] 
        data = os.path.join(pn, 'data', 'dict')
        
        latinalph = ['fr', 'uz', 'en']
        #stackOverflow queries by other users assisted with understanding regex
        
        if ilang in latinalph:
            nstring = re.sub(r"\s*[A-Z]\w*\s*", " ", string).strip()
       
        elif ilang == 'uz':
            return newstring
        elif ilang == 'zh_CN':
            big_cn_gloss = os.path.join(data, 'TSAIWORD (Mandarin).txt')
            cn_gloss = r'' + big_cn_gloss
            big_gloss = open(cn_gloss, 'r', encoding = 'utf-8')
            gloss = str(big_gloss.read().split('\n'))     
            nstring = re.sub(r'(?s:[.,])\Z', ' ', nstring) #removes punctuation
        
            nstring = nstring.replace('  ', ' ').strip() #replaces double spaces with single ones, and if there are still spaces at start/end of string, removes those
            nnstring = nstring.split()
            
            for word in nnstring:
                if word in gloss or d.check(word) == True:
                    continue
                elif word not in gloss:
                    continue
            return newstring
        elif ilang == 'kz':
            return newstring
        
        else:
            return newstring
        
        nstring = re.sub(r'(?s:[.,])\Z', ' ', nstring) #removes punctuation
        
        nstring = nstring.replace('  ', ' ').strip() #replaces double spaces with single ones, and if there are still spaces at start/end of string, removes those
        
        
        d = enchant.Dict('en_GB') #sets the english dictionary so it can be accessed - so if english words pop up they're not marked as errors
    
        '''
        The below code was adapted from a StackOverflow question on how to automate spell-check/auto-corrector. I added the if d.check == True to catch out English words because I am working
        with Russian data, and added the if len > 0 because single characters were getting marked as error-words when they are actually initials - but the rest of the code stayed the same
        '''
        '''
        for err in chkr:
            if d.check(err.word) == True: #sometimes english letters/words might appear in the strings and are raised as errors because it's set to Russian
                continue
            elif len(err.word) > 1: #if the length of the erroneous word is more than one character (i.e. it is not an initial or wayward letter which if corrected here will be made more erroneous
                #print('error word:' + str(err.word)) to check which words are being marked - for me to improve the dictionary/thesaurus further
                if len(err.suggest()) > 0: #if there are suggestion words (i.e. the list is not empty)
                    sug = err.suggest()[0] #take the first word (WILL BE OPTIMISED IN FUTURE USING DIFFLIB SEQUENCE MATCHER RATIOS TO FIND BEST MATCH OF THE WORDS - but this may involve some machine learning)
                    newstring = newstring.replace(str(err.word), str(sug)) #newstring is newstring where the erroneous word is replaced by the suggested substitute
                else: #otherwise if there are no suggested words
                    continue #continue on (likely not a mistake word but one not included in dictionary)
            else: #otherwise continue on
                continue
        '''
        #print('checked ' + str(newstring) + '\n') #test
        return newstring #newstring with replacements of checked words is the return result 

    def doc_spellcheck(self, ilang, word):
        d = enchant.Dict('en_GB')
        sim = dict()
        weird =['{','-', '_', '+', '}', '=', '<', '>', '/', '$', '£', ':', ';', '&', '%', '"',',','.' ]
        pn = os.path.split(os.path.dirname(os.path.abspath(__file__)))[0] 
        data = os.path.join(pn, 'data', 'dict')

        
        if d.check(word) == True:
            return word
        elif word[0] in weird:
            return word
        elif word[-1] in weird:
            word = word[:-2]
        else:
            pass
        
    
            
        if ilang == 'zh_CN':
            big_cn_gloss = os.path.join(data, 'TSAIWORD (Mandarin).txt')
            cn_gloss = r'' + big_cn_gloss
            big_gloss = open(cn_gloss, 'r', encoding = 'utf-8')
            gloss = str(big_gloss.read().split('\n'))     
        elif ilang == 'uz':
            return word
        else:
            return word
            
        if word in gloss:
            #print('found')
            return word
        else:
            return word
            '''
            if r.check(word) == False and word not in gloss:
                suggestions = set(chkr.suggest(word))
                if len(word) > 1:
                    if len(suggestions) == 1:
                        sug = list(suggestions)[0]
                        print('sug1 ' + sug)
                        return str(sug)
                    elif len(suggestions) == 0:
                        print('len0 + ' + word)
                        return word
                    else:
                        print(list(suggestions))
                        for w in list(suggestions):
                            measure = difflib.SequenceMatcher(None, word, w).ratio()
                            sim[measure] = w
                            sug = sim[max(sim.keys())]
                            if sug == None:
                                sug = list(suggestions)[0]
                                print('sug2 ' + sug)
                                return str(sug)
                            else:
                                print('sug3 ' + sug)
                                return str(sug)
                else:
                    return word
            else:
                return word
            '''
    def spell_checker(self, filetype, ilang, word, russgloss):
        
        supported = ['en', 'es', 'fr', 'pt', 'de', 'ru']
        weird =['{','-', '_', '+', '}', '=', '<', '>', '/', '$', '£', ':', ';', '&', '%', '"',',','.' ]
        
        pn = os.path.split(os.path.dirname(os.path.abspath(__file__)))[0] 
        data = os.path.join(pn, 'data', 'dict')
        
        if str(word)[0] in weird and len(str(word)) == 1:
            return word
        elif str(word)[-1] in weird and len(str(word)) > 1:
            word = str(word)[:-2]
        else:
            pass
                
        if ilang in supported:
            spellcheck = sc.SpellChecker(ilang)
            
            if filetype == 'Word' or filetype == 'PDF':
                if ilang == 'ru':
                    russ_words = list(str(s) for s in russgloss.keys())
                    spellcheck.word_frequency.load_words(russ_words)
                    
                    d = enchant.Dict('ru_RU')
                    
                    big_ru_gloss = os.path.join(data, 'ru_dict.txt')
                    ru_gloss = r'' + big_ru_gloss
                    big_gloss = open(ru_gloss, 'r', encoding='utf-8')
                    gloss = str(big_gloss.read()).split('\n')
                    
                elif ilang == 'en':
                    d = enchant.Dict('en_GB')
                    gloss = []
                elif ilang == 'es':
                    d = enchant.Dict('es')
                    gloss = []
                elif ilang == 'fr':
                    d = enchant.Dict('fr')
                    gloss = []
                elif ilang == 'de':
                    d = enchant.Dict('de_DE_frami')
                    gloss = []
                
            
                if word in gloss or d.check(word) == True:
                    #print('found')
                    return word
                elif len(word) > 1:
                    misspelled = spellcheck.unknown(word)
                    if len(misspelled) > 0:
                            #print(spellcheck.candidates(word))
                            return spellcheck.correction(word)
                    else:
                            return word
                else:
                    return word   
                
            elif filetype == 'Excel':
                if ilang == 'ru':
                    russ_words = list(str(s) for s in self.russgloss.keys())
                    spellcheck.word_frequency.load_words(russ_words)
                    
                    #nstring = re.sub(r"\s*[А-Я]\w*\s*", " ", str(word)).strip() #removes all words beginning with capital letters and removes spaces at the start/end
                    nstring = re.sub(r'(?s:(?=(?P<g0>.?«))(?P=g0)(?=(?P<g1>.*?»))(?P=g1).*)\Z', '', str(word)).strip() #uses regex to remove any words within russian parentheses - which includes Central Asian terms and names
                    d = enchant.Dict('ru_RU')
                    
                    big_ru_gloss = os.path.join(data, 'ru_dict.txt')
                    ru_gloss = r'' + big_ru_gloss
                    big_gloss = open(ru_gloss, 'r', encoding='utf-8')
                    gloss = str(big_gloss.read()).split('\n')
                
                else:  
                    nstring = re.sub(r"\s*[A-Z]\w*\s*", " ", word).strip()
                    if ilang == 'en':
                        d = enchant.Dict('en_GB')
                        gloss = []
                    elif ilang == 'es':
                        d = enchant.Dict('es')
                        gloss = []
                    elif ilang == 'fr':
                        d = enchant.Dict('fr')
                        gloss = []
                    elif ilang == 'de':
                        d = enchant.Dict('de_DE_frami')
                        gloss = []
                    
                nstring = re.sub(r'(?s:[.,;!£$&";:?<>%])\Z', '', nstring) #removes punctuation
                nstring = nstring.replace('  ', ' ').strip() #replaces double spaces with single ones, and if there are still spaces at start/end of string, removes those
                split_word = nstring.split(' ')
                for wordd in split_word:
                    try:
                        if wordd[0].isupper(): #if the first word is a capital ignore b/c proper noun
                            continue
                        elif wordd in gloss or d.check(wordd) == True:
                            #print('found')
                            continue
                        elif len(wordd) > 2:
                            if wordd not in gloss or d.check(wordd) == False:
                                checked_word = spellcheck.correction(wordd)
                                word = str(word).replace(wordd, checked_word)
                                continue
                            else:
                                continue
                        else:
                            continue
                    except:
                        continue
                    
                return str(word)
        else:
            if filetype == 'Word' or filetype == 'PDF':
                self.doc_spellcheck(ilang, word)
            elif filetype == 'Excel':
                self.xl_spellcheck(word, ilang)
            else:
                return word
        
            

class ResetGUI:
    def restart(self):
        window = tk.Tk()
        framek = tk.Frame(master=window, width = 150, height = 150, padx=5, pady=5)
        framek.pack(fill = tk.X, side = tk.TOP)
        framei = tk.Frame(master=window, width=150, height = 150, padx = 5, pady=5)
        framei.pack(fill=tk.X, side = tk.TOP)
        rs = tk.Label(master=framek, text = 'Would you like to restart with another data path?')
        rs.pack(fill = tk.X, side = tk.TOP)
        
        def callback(selection):
            if str(selection) == 'Yes':
                window.destroy()
                self.reset()
                
            else:
                framek.forget()
                goodbye = tk.Label(master=framei, text = 'Exiting....Goodbye!')
                goodbye.pack(fill=tk.X, side = tk.TOP)
                window.after(3000, window.destroy)
                sys.exit()
            
        
                
        
        cont = StringVar(framek)
        cont.set('-')
        filetypes = OptionMenu(framek, cont, 'Yes', 'No', command = callback)
        filetypes.pack(fill = tk.X, side = tk.RIGHT)
        window.mainloop()
    def reset(self):
        GUI.start(GUI.start)
        
        
